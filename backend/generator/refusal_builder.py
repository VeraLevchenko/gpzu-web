from docx import Document
from datetime import datetime
from pathlib import Path
import io
import os
import logging
from typing import Dict, Any, Tuple

from database import SessionLocal
from models.application import Application
from models.refusal import Refusal, get_next_refusal_number

logger = logging.getLogger("gpzu-web.refusal_builder")

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates" / "refusal"
ATTACHMENTS_DIR = BASE_DIR / "uploads" / "attachments" / "refusals"
ATTACHMENTS_DIR.mkdir(parents=True, exist_ok=True)

# Маппинг: код причины → файл шаблона
TEMPLATE_FILES = {
    "NO_RIGHTS": "refusal_no_rights.docx",
    "NO_BORDERS": "refusal_no_borders.docx",
    "NOT_IN_CITY": "refusal_not_in_city.docx",
    "OBJECT_NOT_EXISTS": "refusal_object_not_exists.docx",
    "HAS_ACTIVE_GP": "refusal_has_active_gp.docx",
}



REASON_TEXTS = {
    "NO_RIGHTS": (
        "отсутствие у заявителя прав на земельный участок. "
        "В соответствии с подпунктом 1 пункта 6 статьи 57.3 Градостроительного кодекса РФ "
        "градостроительный план земельного участка не подготавливается и не выдается "
        "в случае, если не установлены права на такой земельный участок."
    ),
    "NO_BORDERS": (
        "отсутствие сведений о границах земельного участка в Едином государственном реестре недвижимости. "
        "В соответствии с подпунктом 2 пункта 6 статьи 57.3 Градостроительного кодекса РФ "
        "градостроительный план земельного участка не подготавливается и не выдается "
        "в случае, если в Едином государственном реестре недвижимости отсутствуют сведения "
        "о границах такого земельного участка."
    ),
    "NOT_IN_CITY": (
        "земельный участок не находится на территории городского округа. "
        "В соответствии с пунктом 1 статьи 57.3 Градостроительного кодекса РФ "
        "градостроительный план земельного участка подготавливается применительно к земельным участкам, "
        "расположенным в границах территории, в отношении которой утверждены правила землепользования и застройки."
    ),
    "OBJECT_NOT_EXISTS": (
        "на земельном участке расположен объект капитального строительства. "
        "В соответствии с подпунктом 3 пункта 6 статьи 57.3 Градостроительного кодекса РФ "
        "градостроительный план земельного участка не подготавливается и не выдается "
        "в случае, если на таком земельном участке расположены объекты капитального строительства, "
        "за исключением случаев, предусмотренных частью 1.1 статьи 51.1 настоящего Кодекса."
    ),
    "HAS_ACTIVE_GP": (
        "имеется действующий градостроительный план земельного участка. "
        "Срок действия ранее выданного градостроительного плана не истек."
    ),
}


def get_or_create_application(context: Dict[str, Any], db_session) -> Tuple[int, bool]:
    """
    Находит существующее заявление или создает новое.
    
    Returns:
        Tuple[int, bool]: (application_id, was_created)
        - application_id: ID заявления
        - was_created: True если создано новое, False если найдено существующее
    """
    app_data = context.get('application', {})
    egrn_data = context.get('egrn', {})
    
    app_number = app_data.get('number', '')
    
    existing = db_session.query(Application).filter(Application.number == app_number).first()
    
    if existing:
        logger.info(f"✅ Найдено существующее заявление #{app_number} (ID: {existing.id})")
        existing.status = 'refused'
        db_session.flush()
        return existing.id, False
    
    logger.info(f"📝 Создаем новое заявление #{app_number}")

    app_date = app_data.get('date_formatted') or app_data.get('date', '')
    
    application = Application(
        number=app_number,
        date=app_date,
        applicant=app_data.get('applicant', ''),
        phone=app_data.get('phone', '—'),
        email=app_data.get('email', '—'),
        cadnum=egrn_data.get('cadnum', ''),
        address=egrn_data.get('address', ''),
        area=float(egrn_data.get('area', 0)) if egrn_data.get('area') else None,
        permitted_use=egrn_data.get('vri', ''),
        status='refused'
    )
    
    db_session.add(application)
    db_session.flush()
    
    logger.info(f"✅ Заявление создано (ID: {application.id})")
    return application.id, True


def save_refusal_to_database(context: Dict[str, Any], application_id: int, db_session) -> int:
    """Сохраняет отказ в БД и возвращает ID записи."""
    refusal_data = context.get('refusal', {})
    
    # Проверяем, есть ли уже отказ для этого заявления
    existing_refusal = db_session.query(Refusal).filter(Refusal.application_id == application_id).first()
    
    if existing_refusal:
        logger.warning(f"⚠️ Для заявления ID={application_id} уже есть отказ ID={existing_refusal.id}")
        logger.info(f"🔄 Обновляем существующий отказ вместо создания нового")
        
        # Обновляем существующий отказ
        date_str = refusal_data.get('date', '')
        try:
            out_date_obj = datetime.strptime(date_str, '%d.%m.%Y')
        except ValueError:
            logger.warning(f"⚠️ Неожиданный формат даты: {date_str}, используем текущую дату")
            out_date_obj = datetime.now()
        
        out_date_str = out_date_obj.strftime('%d.%m.%Y')
        out_year = out_date_obj.year
        
        reason_code = refusal_data.get('reason_code', 'NO_RIGHTS')
        reason_text = REASON_TEXTS.get(reason_code, '')
        
        # Обновляем поля
        existing_refusal.out_date = out_date_str
        existing_refusal.out_year = out_year
        existing_refusal.reason_code = reason_code
        existing_refusal.reason_text = reason_text
        
        db_session.flush()
        
        logger.info(f"✅ Отказ обновлен (ID: {existing_refusal.id}, исх. №{existing_refusal.out_number})")
        return existing_refusal.id
    
    # Если отказа нет, создаём новый
    out_number = get_next_refusal_number(db_session)
    
    # Получаем дату и очищаем от лишних символов
    date_str = refusal_data.get('date', '')
    
    # Парсим дату
    try:
        out_date_obj = datetime.strptime(date_str, '%d.%m.%Y')
    except ValueError:
        logger.warning(f"⚠️ Неожиданный формат даты: {date_str}, используем текущую дату")
        out_date_obj = datetime.now()
    
    out_date_str = out_date_obj.strftime('%d.%m.%Y')
    out_year = out_date_obj.year
    
    reason_code = refusal_data.get('reason_code', 'NO_RIGHTS')
    reason_text = REASON_TEXTS.get(reason_code, '')
    
    refusal = Refusal(
        application_id=application_id,
        out_number=out_number,
        out_date=out_date_str,
        out_year=out_year,
        reason_code=reason_code,
        reason_text=reason_text,
        attachment=None
    )
    
    db_session.add(refusal)
    db_session.flush()
    
    logger.info(f"✅ Отказ сохранен в БД (ID: {refusal.id}, исх. №{out_number})")
    return refusal.id


def save_attachment(context: Dict[str, Any], refusal_id: int, docx_bytes: bytes, db_session):
    """Сохраняет вложение отказа на диск и обновляет запись в БД."""
    egrn_data = context.get('egrn', {})
    cadnum = egrn_data.get('cadnum', 'unknown')
    cadnum_safe = cadnum.replace(':', '_')
    
    filename = f"otkaz_{refusal_id}_{cadnum_safe}.docx"
    file_path = ATTACHMENTS_DIR / filename
    
    with open(file_path, 'wb') as f:
        f.write(docx_bytes)
    
    refusal = db_session.query(Refusal).filter(Refusal.id == refusal_id).first()
    if refusal:
        refusal.attachment = str(file_path)
        db_session.flush()
    
    logger.info(f"💾 Вложение сохранено: {file_path}")


def write_refusal_to_excel_journal(context: Dict[str, Any]):
    """Дублирование в Excel журнал (опционально, для переходного периода)."""
    pass


def build_refusal_doc(context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Генерирует документ отказа в выдаче ГПЗУ и сохраняет в БД.
    
    Returns:
        Dict с ключами:
        - document: bytes документа
        - application_created: bool (True если заявление создано, False если найдено существующее)
        - refusal_id: int (ID записи отказа в БД)
    """
    db = SessionLocal()
    application_created = False
    refusal_id = None
    
    try:
        application_id, was_created = get_or_create_application(context, db)
        application_created = was_created
        
        refusal_id = save_refusal_to_database(context, application_id, db)
        
        # Получаем код причины отказа
        refusal_data = context.get('refusal', {})
        reason_code = refusal_data.get('reason_code', 'NO_RIGHTS')
        
        # Выбираем нужный шаблон
        template_filename = TEMPLATE_FILES.get(reason_code)
        if not template_filename:
            logger.error(f"❌ Неизвестная причина отказа: {reason_code}")
            raise Exception(f"Неизвестная причина отказа: {reason_code}")
        
        template_path = TEMPLATES_DIR / template_filename
        
        # Проверяем наличие шаблона
        if not template_path.exists():
            logger.error(f"❌ Шаблон не найден: {template_path}")
            raise Exception(f"Шаблон отказа не найден: {template_filename}")
        
        logger.info(f"📄 Используется шаблон: {template_filename}")
        doc = Document(str(template_path))
        
        app_data = context.get('application', {})
        egrn_data = context.get('egrn', {})
        
        app_number = app_data.get('number', '')
        app_date = app_data.get('date', '')
        applicant = app_data.get('applicant', '')
        
        cadnum = egrn_data.get('cadnum', '')
        address = egrn_data.get('address', '')
        
        out_number = db.query(Refusal).filter(Refusal.id == refusal_id).first().out_number
        out_date = refusal_data.get('date', '')
        
        reason_text = REASON_TEXTS.get(reason_code, '')
        specialist = context.get('specialist', '—')

        # Получаем телефон и email из данных заявления
        phone = app_data.get('phone', '—')
        email = app_data.get('email', '—')

        replacements = {
            '{{OUT_NUMBER}}': str(out_number),
            '{{OUT_NUM}}': str(out_number),  # Добавлен альтернативный плейсхолдер
            '{{OUT_DATE}}': out_date,
            '{{APP_NUMBER}}': app_number,
            '{{APP_DATE}}': app_date,
            '{{APPLICANT}}': applicant,
            '{{PHONE}}': phone,  # Добавлен телефон
            '{{EMAIL}}': email,  # Добавлен email
            '{{CADNUM}}': cadnum,
            '{{ADDRESS}}': address,
            '{{REASON}}': reason_text,
            '{{SPECIALIST}}': specialist,
        }
        
        # Функция замены текста в параграфе
        def replace_in_paragraph(paragraph, replacements):
            """Заменяет плейсхолдеры в параграфе."""
            # Сначала собираем весь текст параграфа
            full_text = paragraph.text
            
            # Проверяем, есть ли что заменять
            has_placeholder = any(key in full_text for key in replacements.keys())
            if not has_placeholder:
                return
            
            # Выполняем замену
            new_text = full_text
            for key, value in replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, str(value))
            
            # Если текст изменился, обновляем параграф
            if new_text != full_text:
                # Удаляем все runs
                for run in paragraph.runs:
                    run.text = ''
                
                # Добавляем новый текст
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        
        # Замена в параграфах документа
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        
        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

        # Замена в колонтитулах (headers и footers)
        logger.info("🔍 Начинаем замену в колонтитулах...")
        for section_idx, section in enumerate(doc.sections):
            logger.info(f"  Секция {section_idx}:")
            
            # Замена в верхних колонтитулах
            header_count = 0
            for para in section.header.paragraphs:
                if para.text.strip():
                    logger.info(f"    Header paragraph: '{para.text[:50]}'")
                    replace_in_paragraph(para, replacements)
                    header_count += 1
            
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                logger.info(f"    Header table cell: '{para.text[:50]}'")
                                replace_in_paragraph(para, replacements)
                                header_count += 1
            
            logger.info(f"    Header: обработано {header_count} элементов")
            
            # Замена в нижних колонтитулах
            footer_count = 0
            for para in section.footer.paragraphs:
                if para.text.strip():
                    logger.info(f"    Footer paragraph: '{para.text[:50]}'")
                    if '{{SPECIALIST}}' in para.text:
                        logger.info(f"      ⚠️ Найден {{{{SPECIALIST}}}} в тексте!")
                    replace_in_paragraph(para, replacements)
                    logger.info(f"      После замены: '{para.text[:50]}'")
                    footer_count += 1
            
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                logger.info(f"    Footer table cell: '{para.text[:50]}'")
                                replace_in_paragraph(para, replacements)
                                footer_count += 1
            
            logger.info(f"    Footer: обработано {footer_count} элементов")

        logger.info("✅ Замена в колонтитулах завершена")
        
        # Сохранение документа
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_bytes = docx_buffer.read()
        
        save_attachment(context, refusal_id, docx_bytes, db)
        
        write_refusal_to_excel_journal(context)
        
        db.commit()
        
        logger.info(f"✅ Отказ успешно сформирован (ID: {refusal_id}, заявление {'создано' if application_created else 'найдено'})")
        
        return {
            'document': docx_bytes,
            'application_created': application_created,
            'refusal_id': refusal_id
        }
        
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Ошибка генерации отказа: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise
    finally:
        db.close()