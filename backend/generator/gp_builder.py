# backend/generator/gp_builder.py
"""
Генератор градостроительного плана земельного участка (ГПЗУ).
ИСПРАВЛЕННАЯ ВЕРСИЯ на основе оригинала + поддержка площадей ЗОУИТ
"""

import os
import logging
from copy import deepcopy
from pathlib import Path
from typing import Dict, Any, Optional, List
import re

from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Cm, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.table import Table, _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("gpzu-web.gp_builder")

# ----------------- Таблица координат ----------------- #

COL_W = [Cm(4.50), Cm(6.69), Cm(6.69)]
MARKER_COORDS = "[[COORDS_TABLE]]"

# ----------------- ФУНКЦИЯ ФОРМАТИРОВАНИЯ ПЛОЩАДЕЙ ----------------- #

def format_area(area_value) -> str:
    """
    Форматирует площадь в русском формате: 1024.46 → "1 024,46"
    """
    if area_value is None:
        return ""
    
    try:
        if isinstance(area_value, str):
            area_value = area_value.strip().replace(" ", "").replace(",", ".")
            if not area_value:
                return ""
        
        num_area = float(area_value)
        
        if num_area <= 0:
            return ""
        
        # Форматируем число: пробел как разделитель тысяч, запятая как десятичный разделитель
        formatted = f"{num_area:,.2f}"  # Получаем "1,024.46"
        formatted = formatted.replace(",", " ")  # тысячи: запятая → пробел
        formatted = formatted.replace(".", ",")  # десятичные: точка → запятая
        
        return formatted
        
    except (ValueError, TypeError):
        return ""

# ----------------- ОБНОВЛЕННОЕ СОПОСТАВЛЕНИЕ ЗОУИТ ----------------- #

# Словарь сопоставления ключевых слов с РЕАЛЬНО СУЩЕСТВУЮЩИМИ файлами ЗОУИТ
ZOUIT_MAPPING = {
    # Водоохранная зона (55)
    "55_vodoohr_pribr_bereg.docx": [
        "водоохран", "прибрежн", "защитн полос", "водных объект", "водоохранная зона", 
        "водоохранных зон", "прибрежная защитная полоса", "прибрежных защитных полос", 
        "водный объект", "водные объекты", "береговая полоса"
    ],
    
    # Санитарно-защитные зоны (56)
    "56_sanzona.docx": [
        "санитар", "санитарно-защит", "санитарно защит", "сзз", 
        "санитарная зона", "санитарно-защитная зона", "санитарные зоны"
    ],
    
    # Электросетевое хозяйство (57)
    "57_electro.docx": [
        "охранная зона объектов электросетевого хозяйства", "ВЛ", "КЛ", "ВЛ-КЛ",
        "охранная зона вл", "охранная зона кл", "электроэнергетики",
        "сооружение линейное электротехническое", "электротехническое сооружение",
        "воздушной линии электропередачи", "воздушная линия электропередач",
        "электропередач", "вл-кл", "воздушная линия", "кабельная линия",
        "линия электропередач", "электросетевое хозяйство", "электросетевого хозяйства",
        "электрические сети", "электросети", "линии электропередач", "электросетей"
    ],
    
    # Газораспределительные сети (57)
    "57_gazoraspredelitelnyh_setey.docx": [
        "газ", "газораспределительн", "газопровод", "газоснабжен", 
        "газораспределительная сеть", "газораспределительных сетей",
        "газопроводов", "газоснабжения", "магистральный газопровод", "газовые сети"
    ],
    
    # Линии и сооружения связи (57)
    "57_linii_i_sooruzheniy_svyazi.docx": [
        "связь", "линии связи", "кабельн связ", "сооружения связи",
        "линий связи", "кабели связи", "кабельные линии связи", "сооружений связи"
    ],
    
    # Тепловые сети (57)
    "57_teplovyh_setey.docx": [
        "тепло", "теплотрасс", "теплосет", "теплоснабжен",
        "тепловые сети", "тепловых сетей", "теплоснабжения",
        "теплотрассы", "теплосетей", "теплопровод", "тепловых сетей"
    ],
    
    # Охрана объектов культурного наследия (58)
    "58_ohrany_obektov_kulturnogo_naslediya.docx": [
        "культурн наследи", "окн", "объект культурного наследия", "охранная зона окн",
        "зона охраны объектов культурного наследия", "памятник", "памятники",
        "культурного наследия", "объекты культурного наследия"
    ],
    
    # Затопления и подтопления территорий (59)
    "59_zatopleniya_i_podtopleniya_territoriy.docx": [
        "затоплен", "умеренного подтопления", "слабого подтопления", "сильного подтопления",  "подтоплен", "зона затопления", "зона подтопления",
        "затопления территории", "подтопления территории", "паводк", "наводнен"
    ],
    
    # Охранные зоны пунктов наблюдения окружающей среды (60)
    "60_ohrannye_zony_pn_nablyudeniya_okr_sredy.docx": [
        "наблюдени", "окружающ", "среды", "экологическ", "мониторинг",
        "пункт наблюдения", "охранная зона пункта наблюдения"
    ],
    
    # Охранные зоны геодезической сети (61)
    "61_ohrannye_zony_geodezicheskoy_seti.docx": [
        "геодезическ", "геодезическая сеть", "геодезических пунктов", "триангуляци",
        "нивелирн", "охранная зона геодезических", "геодезический пункт"
    ],
    
    # Зоны магистральных трубопроводов (62)
    "62_zony_magistralnyh_truboprovodov.docx": [
        "магистральн", "трубопровод", "нефтепровод", "нефтепродуктопровод", 
        "магистральный трубопровод", "магистральных трубопроводов",
        "нефтепроводов", "нефтепродуктопроводов", "газопровод магистральный"
    ],
    
    # Зоны радиотехнического объекта (63)
    "63_zony_radiotehnicheskogo_obekta.docx": [
        "радиотехническ", "радио", "антенн", "передающ", "радиолокаци",
        "радиотехнический объект", "радиотехнических объектов", "радиооборудован"
    ],
    
    # Приаэродромная территория (64)
    "64_aeroport_full.docx": [
        "приаэродром", "аэродром", "аэропорт", "приаэродромная территория"
    ],
    "64_aeroport_podzona1.docx": [
        "перв подзон", "первая подзона", "1 подзона", "первой подзоны", "подзона 1"
    ],
    "64_aeroport_podzona2.docx": [
        "втор подзон", "вторая подзона", "2 подзона", "второй подзоны", "подзона 2"
    ],
    "64_aeroport_podzona3.docx": [
        "трет подзон", "третья подзона", "3 подзона", "третьей подзоны", "подзона 3"
    ],
    "64_aeroport_podzona4.docx": [
        "четверт подзон", "четвертая подзона", "4 подзона", "четвертой подзоны", "подзона 4"
    ],
    "64_aeroport_podzona5.docx": [
        "Пятая", "Пятая подзона", "пятая подзона", "5 подзона", "пятой подзоны", "подзона 5"
    ],
    "64_aeroport_podzona6.docx": [
        "шест подзон", "шестая подзона","Шестая подзона", "6 подзона", "шестой подзоны", "подзона 6"
    ],
    "64_aeroport_podzona7.docx": [
        "седьм подзон", "седьмая подзона", "7 подзона", "седьмой подзоны", "подзона 7"
    ],
    
    # Санитарная зона источников водоснабжения (65)
    "65_sanitarnaya_zona_istochnikov_vodosnabzheniya.docx": [
        "санитарная зона источник", "источник водоснабжения", "водозабор",
        "скважин", "колодц", "источники водоснабжения", "водозаборных сооружений"
    ],
    
    # Охранная зона железных дорог (66)
    "66_ohrannaya_zona_jeleznodorog.docx": [
        "железн дорог", "железнодорожн", "жд ", " жд", "железная дорога",
        "железных дорог", "железнодорожного транспорта", "охранная зона железных дорог"
    ]
}

# Правила высокого приоритета: если подстрока встречается в названии,
# сразу выбираем указанный файл, без сложного скоринга
HIGH_PRIORITY_ZOUIT_RULES = [
    # Водоохранные зоны
    ("водоохранная зона", "55_vodoohr_pribr_bereg.docx"),
    ("прибрежная защитная полоса", "55_vodoohr_pribr_bereg.docx"),

    # Санитарно-защитные зоны
    ("санитарно-защитная зона", "56_sanzona.docx"),
    ("санитарно защитная зона", "56_sanzona.docx"),
    ("санитарно-защитная зона для полигона", "56_sanzona.docx"),

    # Электросетевое хозяйство
    ("объектов электросетевого хозяйства", "57_electro.docx"),
    ("сооружение линейное электротехническое", "57_electro.docx"),
    ("воздушная линия электропередачи", "57_electro.docx"),
    ("кабельная линия электропередачи", "57_electro.docx"),
    ("вл-", "57_electro.docx"),   # ВЛ-0,4; ВЛ-6 и т.п.
    ("вли-0,4", "57_electro.docx"),

    # Газораспределительные сети
    ("газопровод", "57_gazoraspredelitelnyh_setey.docx"),
    ("газораспределительн", "57_gazoraspredelitelnyh_setey.docx"),
    ("газоснабжен", "57_gazoraspredelitelnyh_setey.docx"),

    # Тепловые сети
    ("тепловые сети", "57_teplovyh_setey.docx"),
    ("теплотрасс", "57_teplovyh_setey.docx"),
    ("теплопровод", "57_teplovyh_setey.docx"),

    # Линии и сооружения связи
    ("линии связи", "57_linii_i_sooruzheniy_svyazi.docx"),
    ("кабельные линии связи", "57_linii_i_sooruzheniy_svyazi.docx"),
    ("сооружения связи", "57_linii_i_sooruzheniy_svyazi.docx"),

    # Объекты культурного наследия
    ("объектов культурного наследия", "58_ohrany_obektov_kulturnogo_naslediya.docx"),
    ("памятник архитектуры", "58_ohrany_obektov_kulturnogo_naslediya.docx"),

    # Затопление / подтопление
    ("зона затопления территорий", "59_zatopleniya_i_podtopleniya_territoriy.docx"),
    ("подтопления территории", "59_zatopleniya_i_podtopleniya_territoriy.docx"),

    # Пункты наблюдения
    ("пункта наблюдения за состоянием окружающей среды", "60_ohrannye_zony_pn_nablyudeniya_okr_sredy.docx"),
    ("экологического мониторинга", "60_ohrannye_zony_pn_nablyudeniya_okr_sredy.docx"),

    # Геодезическая сеть
    ("геодезическ", "61_ohrannye_zony_geodezicheskoy_seti.docx"),

    # Магистральные трубопроводы
    ("магистральный трубопровод", "62_zony_magistralnyh_truboprovodov.docx"),
    ("нефтепровод", "62_zony_magistralnyh_truboprovodov.docx"),
    ("нефтепродуктопровод", "62_zony_magistralnyh_truboprovodov.docx"),

    # Радиотехнические объекты
    ("радиотехническ", "63_zony_radiotehnicheskogo_obekta.docx"),
    ("радиолокационн", "63_zony_radiotehnicheskogo_obekta.docx"),

    # Приаэродромная территория – конкретные подзоны
    ("первая подзона приаэродромной территории", "64_aeroport_podzona1.docx"),
    ("вторая подзона приаэродромной территории", "64_aeroport_podzona2.docx"),
    ("третья подзона приаэродромной территории", "64_aeroport_podzona3.docx"),
    ("четвертая подзона приаэродромной территории", "64_aeroport_podzona4.docx"),
    ("пятая подзона приаэродромной территории", "64_aeroport_podzona5.docx"),
    ("шестая подзона приаэродромной территории", "64_aeroport_podzona6.docx"),
    ("седьмая подзона приаэродромной территории", "64_aeroport_podzona7.docx"),

    # Обобщённая приаэродромная территория (без указания подзоны)
    ("приаэродромной территории аэродрома", "64_aeroport_full.docx"),
    ("приаэродромная территория аэродрома", "64_aeroport_full.docx"),

    # Зона санитарной охраны источников водоснабжения
    ("санитарная зона источников водоснабжения", "65_sanitarnaya_zona_istochnikov_vodosnabzheniya.docx"),
    ("зона санитарной охраны водозабора", "65_sanitarnaya_zona_istochnikov_vodosnabzheniya.docx"),

    # Железные дороги
    ("охранная зона железной дороги", "66_ohrannaya_zona_jeleznodorog.docx"),
    ("железнодорожного транспорта", "66_ohrannaya_zona_jeleznodorog.docx"),
]

# Специальные случаи по реестровым номерам
SPECIAL_REGISTRY_MAPPING = {
    "42:00-6.1695": "64_aeroport_full.docx"  # Приаэродромная территория целиком
}


def _center_cell(cell: _Cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for par in cell.paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_table_layout(tbl: Table):
    """Фиксированная ширина и выравнивание таблицы координат."""
    try:
        tbl.autofit = False
    except Exception:
        pass
    try:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(COL_W):
                try:
                    cell.width = COL_W[i]
                except Exception:
                    pass
            _center_cell(cell)


def _fmt_coord(v: Optional[str]) -> str:
    """Формат числа: без пробелов, с запятой как разделителем."""
    return (v or "").strip().replace(" ", "").replace(".", ",")


def _iter_all_paragraphs(doc: Document):
    """Итерация по всем параграфам, включая те, что внутри таблиц."""
    for p in doc.paragraphs:
        yield p

    def walk_cell(cell: _Cell):
        for p in cell.paragraphs:
            yield p
        for t in cell.tables:
            for r in t.rows:
                for c in r.cells:
                    yield from walk_cell(c)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                yield from walk_cell(c)


def _find_paragraph_with_text(doc: Document, marker: str):
    """Найти первый параграф, содержащий данный текст."""
    for p in _iter_all_paragraphs(doc):
        if p.text and marker in p.text:
            return p
    return None


def _replace_paragraph_with_table(anchor_paragraph, table: Table):
    """Вставить таблицу сразу после параграфа и удалить сам параграф."""
    anchor_elm = anchor_paragraph._element
    parent = anchor_elm.getparent()
    parent.insert(parent.index(anchor_elm) + 1, table._tbl)
    parent.remove(anchor_elm)


class GPBuilder:
    """
    ИСПРАВЛЕННЫЙ генератор градостроительного плана земельного участка (ГПЗУ).
    На основе оригинала + поддержка площадей ЗОУИТ.
    """

    def __init__(self, template_path: str, data_dir: Optional[str] = None):
        self.template_path = str(template_path)

        base_dir = Path(__file__).resolve().parent.parent
        if data_dir is None:
            self.data_dir = base_dir / "data"
        else:
            self.data_dir = Path(data_dir)

        self.tz_dir = self.data_dir / "tz_reglament"
        self.zouit_dir = self.data_dir / "zouit_reglament"
        self.ago_reglament_dir = self.data_dir / "ago_reglament"

        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Шаблон не найден: {self.template_path}")

        logger.info(f"GPBuilder: шаблон: {self.template_path}")
        logger.info(f"GPBuilder: data_dir: {self.data_dir}")
        logger.info(f"GPBuilder: tz_dir: {self.tz_dir}")
        logger.info(f"GPBuilder: zouit_dir: {self.zouit_dir}")


    def _normalize_zouit_name(self, name: str) -> str:
        """
        Нормализация названия ЗОУИТ для сопоставления.
        """
        if not name:
            return ""

        # 1) нижний регистр
        normalized = name.strip().lower()

        # 2) заменяем латинские буквы на кириллицу, если они "похожи"
        latin_to_cyr = str.maketrans({
            "a": "а",  # лат. a -> рус. а
            "b": "в",
            "c": "с",
            "e": "е",
            "h": "н",
            "k": "к",
            "m": "м",
            "o": "о",
            "p": "р",
            "t": "т",
            "x": "х",
            "y": "у",
        })
        normalized = normalized.translate(latin_to_cyr)

        # 3) убираем кавычки и дублирующиеся пробелы
        normalized = re.sub(r'[«»"\'""]+', " ", normalized)
        normalized = re.sub(r"\s+", " ", normalized).strip()

        # 4) убираем длинные сервисные префиксы
        prefixes_to_remove = [
            "зона с особыми условиями использования территории (охранная зона)",
            "зона с особыми условиями использования территории",
            "зона с особыми условиями использования",
            "зона с особыми условиями",
            "зона с особыми условиями использования территории (зоуит)",
            "зоуит",
            "охранная зона",
        ]

        for prefix in prefixes_to_remove:
            if normalized.startswith(prefix):
                normalized = normalized[len(prefix):].strip()
                break

        return normalized


    def _calculate_match_score(self, normalized_name: str, keywords: List[str]) -> float:
        """
        Вычисляет оценку совпадения названия ЗОУИТ с ключевыми словами.
        """
        if not normalized_name or not keywords:
            return 0.0
        
        total_score = 0.0
        matches = 0
        
        for keyword in keywords:
            keyword_lower = keyword.lower()
            
            # Точное совпадение - максимальный балл
            if keyword_lower == normalized_name:
                return 1.0
            
            # Полное вхождение ключевого слова
            if keyword_lower in normalized_name:
                # Бонус за точность совпадения
                accuracy = len(keyword_lower) / len(normalized_name)
                total_score += 0.8 * accuracy
                matches += 1
                continue
            
            # Частичное совпадение слов
            keyword_words = keyword_lower.split()
            name_words = normalized_name.split()
            
            word_matches = 0
            for kw_word in keyword_words:
                for name_word in name_words:
                    # Точное совпадение слова
                    if kw_word == name_word:
                        word_matches += 1
                    # Вхождение слова (для слов длиннее 3 символов)
                    elif len(kw_word) > 3 and (kw_word in name_word or name_word in kw_word):
                        word_matches += 0.5
            
            if word_matches > 0:
                word_score = word_matches / max(len(keyword_words), len(name_words))
                total_score += 0.6 * word_score
                matches += 1
        
        # Нормализуем итоговую оценку
        if matches > 0:
            return min(total_score / matches, 1.0)
        else:
            return 0.0


    def get_zouit_block_filename(self, zouit_name: str) -> Optional[str]:
        """
        Определение имени файла блока ЗОУИТ по наименованию.
        """
        if not zouit_name:
            return None

        # Сначала приводим к нижнему регистру и заменяем латиницу на кириллицу,
        # чтобы high-priority правила тоже работали на "кривых" строках
        raw_lower = zouit_name.strip().lower()

        latin_to_cyr = str.maketrans({
            "a": "а",
            "b": "в",
            "c": "с",
            "e": "е",
            "h": "н",
            "k": "к",
            "m": "м",
            "o": "о",
            "p": "р",
            "t": "т",
            "x": "х",
            "y": "у",
        })
        raw_lower = raw_lower.translate(latin_to_cyr)

        # 1) Пытаемся применить high-priority правила
        for pattern, filename in HIGH_PRIORITY_ZOUIT_RULES:
            if pattern in raw_lower:
                logger.info(
                    f"🎯 HIGH-PRIORITY правило: '{zouit_name}' -> {filename} (паттерн: '{pattern}')"
                )
                return filename

        # 2) Если явных правил нет — используем нормализованное имя и скоринговый поиск
        name_normalized = self._normalize_zouit_name(zouit_name)
        logger.debug(f"Поиск файла для ЗОУИТ: '{zouit_name}' -> '{name_normalized}'")

        best_match = None
        best_score = 0.0

        for filename, keywords in ZOUIT_MAPPING.items():
            score = self._calculate_match_score(name_normalized, keywords)
            if score > best_score:
                best_score = score
                best_match = filename

        if best_match and best_score > 0:
            logger.info(
                f"✅ Найдено совпадение для ЗОУИТ '{zouit_name}': {best_match} (score: {best_score:.2f})"
            )
            return best_match

        logger.warning(f"❌ Не найден файл для ЗОУИТ: '{zouit_name}'")
        return None


    def get_zouit_file(self, zouit_name: str) -> Optional[str]:
        """Совместимость с test_gp_builder.py"""
        return self.get_zouit_block_filename(zouit_name)


    def get_zouit_block_path(self, zouit: Dict[str, Any]) -> Optional[Path]:
        """
        ОБНОВЛЕННАЯ функция получения пути к файлу блока ЗОУИТ.
        """
        name = zouit.get("name") or ""
        registry_number = (zouit.get("registry_number") or "").strip()

        filename: Optional[str]

        # Специальные случаи по реестровому номеру
        if registry_number in SPECIAL_REGISTRY_MAPPING:
            filename = SPECIAL_REGISTRY_MAPPING[registry_number]
            logger.info(f"🎯 Специальный случай для реестрового номера {registry_number}: {filename}")
        else:
            # Используем улучшенную функцию сопоставления
            filename = self.get_zouit_block_filename(name)

        if not filename:
            return None

        path = self.zouit_dir / filename
        if not path.exists():
            logger.warning(f"❌ Файл блока ЗОУИТ не найден: {path}")
            return None
        
        logger.info(f"✅ Найден файл блока ЗОУИТ: {path}")
        return path


    def load_zone_block(self, zone_code: str, block_type: str) -> Optional[Document]:
        """Загружает Word-блок для территориальной зоны."""
        if block_type == "vri":
            filename = f"{zone_code}_vri.docx"
        else:
            filename = f"{zone_code}.docx"

        filepath = self.tz_dir / filename
        if not filepath.exists():
            logger.warning(f"Файл блока зоны не найден: {filepath}")
            return None

        logger.info(f"Загружен блок зоны: {filepath}")
        return Document(str(filepath))


    def load_zouit_block(self, zouit_name: str) -> Optional[Document]:
        """СТАРЫЙ интерфейс (для совместимости)."""
        filename = self.get_zouit_block_filename(zouit_name)
        if not filename:
            return None

        filepath = self.zouit_dir / filename
        if not filepath.exists():
            logger.warning(f"Файл блока ЗОУИТ не найден: {filepath}")
            return None

        logger.info(f"Загружен блок ЗОУИТ (legacy): {filepath}")
        return Document(str(filepath))


    def prepare_context(self, gp_data: Dict[str, Any]) -> Dict[str, Any]:
        """Подготавливает контекст для шаблона."""
        context = dict(gp_data)

        # === ИСПРАВЛЕНИЕ: Форматирование даты заявления === #
        application = gp_data.get("application") or {}
        app_date = application.get("date")  # Может быть строка "2025-11-21" или "«21» ноября 2025 г."

        if app_date:
            try:
                from datetime import datetime
                date_str = str(app_date).strip()
                
                # Убираем время если есть (например "2025-11-21 00:00:00" -> "2025-11-21")
                if " " in date_str:
                    date_str = date_str.split()[0]
                
                # Пробуем распарсить разные форматы
                dt = None
                
                # Попытка 1: ISO формат YYYY-MM-DD
                for fmt in ["%Y-%m-%d", "%d.%m.%Y", "%Y/%m/%d", "%d/%m/%Y"]:
                    try:
                        dt = datetime.strptime(date_str, fmt)
                        break
                    except ValueError:
                        continue
                
                # Если не получилось распарсить стандартные форматы,
                # пробуем формат «21» ноября 2025 г.
                if not dt and "«" in app_date and "»" in app_date:
                    try:
                        # Словарь месяцев
                        months = {
                            "января": 1, "февраля": 2, "марта": 3, "апреля": 4,
                            "мая": 5, "июня": 6, "июля": 7, "августа": 8,
                            "сентября": 9, "октября": 10, "ноября": 11, "декабря": 12,
                        }
                        
                        # Извлекаем день из «21»
                        day_part = app_date.split("«", 1)[1].split("»", 1)[0].strip()
                        day = int(day_part)
                        
                        # Извлекаем остальную часть: "ноября 2025 г."
                        rest = app_date.split("»", 1)[1].strip()
                        rest = rest.replace("г.", "").replace("г", "").strip()
                        parts = rest.split()
                        
                        if len(parts) >= 2:
                            month_name = parts[0].lower()
                            year = int(parts[1])
                            month = months.get(month_name)
                            
                            if month and day and year:
                                dt = datetime(year, month, day)
                    
                    except Exception as ex:
                        logger.warning(f"Не удалось распарсить дату '{app_date}': {ex}")
                
                # Форматируем в DD.MM.YYYY
                if dt:
                    context["application_date_formatted"] = dt.strftime("%d.%m.%Y")
                    logger.info(f"✅ Дата заявления отформатирована: {context['application_date_formatted']}")
                else:
                    # Если не удалось распарсить - оставляем как есть
                    context["application_date_formatted"] = str(app_date)
                    logger.warning(f"⚠️ Дата заявления не распознана, используется как есть: {app_date}")
                    
            except Exception as ex:
                logger.error(f"❌ Ошибка форматирования даты: {ex}")
                context["application_date_formatted"] = str(app_date) if app_date else ""
        else:
            context["application_date_formatted"] = ""
            logger.warning("⚠️ Дата заявления отсутствует")

        # НОВОЕ: Информация о районе
        district = gp_data.get("district") or {}
        district_name = district.get("name") or ""

        # Устанавливаем переменную для шаблона
        context["district_name"] = district_name if district_name else "Район не определён"

        # АГО — архитектурно-градостроительный облик (п.12)
        # Индекс в слое: "АГО-01", "АГО-02" (или "АГО-1", "АГО-2")
        ago_index = gp_data.get("ago_index") or ""
        ago_num = self._ago_number(ago_index)  # "1" или "2" или ""
        if ago_num == "1":
            context["ago_text"] = (
                "Земельный участок расположен в границах территории регулирования "
                "архитектурно-градостроительного облика – 1 (АГО-1)"
            )
        elif ago_num == "2":
            context["ago_text"] = (
                "Земельный участок расположен в границах территории регулирования "
                "архитектурно-градостроительного облика – 2 (АГО-2)"
            )
        else:
            context["ago_text"] = (
                "требования к архитектурно-градостроительному облику объекта "
                "капитального строительства не установлены"
            )
        context["ago_index"] = ago_index
        context["INSERT_AGO_BLOCK"] = "{{INSERT_AGO_BLOCK}}"

        # Объекты капитального строительства
        capital_objects = gp_data.get("capital_objects") or []

        if capital_objects:
            # Подсчитываем количество объектов
            count = len(capital_objects)
            
            # Правильное склонение слова "единица"
            if count == 1:
                unit_word = "единица"
            elif 2 <= count <= 4:
                unit_word = "единицы"
            else:
                unit_word = "единиц"
            
            # Формируем текст
            context["capital_objects_text"] = (
                f"в границах земельного участка расположены объекты капитального строительства. "
                f"Количество объектов {count} {unit_word}. "
                f"Объекты отображаются на чертеже(ах) градостроительного плана под порядковыми номерами. "
                f"Описание объектов капитального строительства приводится в подразделе 3.1 "
                f'"Объекты капитального строительства" или подразделе 3.2 '
                f'"Объекты, включенные в единый государственный реестр объектов культурного наследия '
                f'(памятников истории и культуры) народов Российской Федерации" раздела 3'
            )
            context["capital_objects_count"] = count
        else:
            context["capital_objects_text"] = "Не имеется"
            context["capital_objects_count"] = 0

        # ЗОУИТ в удобном виде для таблицы (раздел 6)
        zouit_raw = gp_data.get("zouit") or []
        formatted: List[Dict[str, str]] = []
        for z in zouit_raw:
            name = z.get("name") or ""
            registry_number = z.get("registry_number") or ""
            document = z.get("document") or ""
            restrictions = z.get("restrictions") or ""

            title = name
            if registry_number:
                title += f" ({registry_number})"

            formatted.append({
                "title": title,
                "document": document,
                "restrictions": restrictions,
            })

        context["zouit_formatted"] = formatted

        # Сохранить маркеры вставки блоков после рендера Jinja
        context["INSERT_ZONE_VRI"] = "{{INSERT_ZONE_VRI}}"
        context["INSERT_ZONE_PARAMS"] = "{{INSERT_ZONE_PARAMS}}"
        context["INSERT_ZOUIT_BLOCKS"] = "{{INSERT_ZOUIT_BLOCKS}}"

        return context


    def insert_block_at_marker(self, doc: Document, marker: str, block_doc: Document) -> None:
        """Вставляет содержимое блока в документ на место маркера."""
        marker_para = None
        for para in doc.paragraphs:
            if marker in para.text:
                marker_para = para
                break

        if marker_para is None:
            logger.warning(f"Маркер {marker!r} не найден в документе")
            return

        # Удаляем сам маркер из текста абзаца
        marker_para.text = marker_para.text.replace(marker, "").strip()

        body = marker_para._p.getparent()
        idx = body.index(marker_para._p)

        # Клонируем элементы из блока (параграфы и таблицы) и вставляем
        elements = [deepcopy(el) for el in block_doc.element.body]
        for el in reversed(elements):
            body.insert(idx + 1, el)


    @staticmethod
    def _ago_number(ago_index: str) -> str:
        """
        Извлекает номер АГО из индекса.
        "АГО-01" → "1", "АГО-1" → "1", "АГО-02" → "2", "" → ""
        """
        if not ago_index:
            return ""
        parts = ago_index.split("-")
        if len(parts) >= 2:
            try:
                return str(int(parts[-1]))
            except ValueError:
                pass
        return ""


    def insert_ago_block(self, doc: Document, ago_index: Optional[str]) -> None:
        """
        Вставляет текст регламента АГО в п.12 на место маркера {{INSERT_AGO_BLOCK}}.
        Если ago_index не задан — просто удаляет маркер.
        """
        marker = "{{INSERT_AGO_BLOCK}}"

        ago_num = self._ago_number(ago_index or "")

        if not ago_num:
            # Удаляем маркер без вставки блока
            for para in doc.paragraphs:
                if marker in para.text:
                    para.text = para.text.replace(marker, "").strip()
                    logger.info("Маркер {{INSERT_AGO_BLOCK}} удалён (АГО не найдено)")
                    break
            return

        filename = f"ago-{ago_num}.docx"  # "ago-1.docx" или "ago-2.docx"
        ago_path = self.ago_reglament_dir / filename

        if not ago_path.exists():
            logger.warning(f"Файл регламента АГО не найден: {ago_path}")
            for para in doc.paragraphs:
                if marker in para.text:
                    para.text = para.text.replace(marker, "").strip()
                    break
            return

        block_doc = Document(str(ago_path))
        self.insert_block_at_marker(doc, marker, block_doc)
        logger.info(f"Вставлен блок регламента АГО: {ago_path.name}")


    def fill_zouit_table(self, doc: Document, zouit_list: List[Dict[str, Any]]) -> None:
        """
        Заполняет таблицу ЗОУИТ в разделе 6 градплана.
        """
        if not zouit_list:
            logger.info("ЗОУИТ отсутствуют, таблица ЗОУИТ не заполняется")
            return

        # Ищем нужную таблицу по тексту в первой ячейке
        target_table: Optional[Table] = None
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                first_cell_text = (table.rows[0].cells[0].text or "").lower()
                if "наименование зоны с особыми условиями использования" in first_cell_text:
                    target_table = table
                    break

        if target_table is None:
            logger.warning("Таблица ЗОУИТ в документе не найдена")
            return

        # Пробуем навесить стиль с границами
        try:
            target_table.style = "Table Grid"
        except Exception:
            # если стиля нет, используем тот, что в шаблоне
            pass

        rows = target_table.rows

        # В текущем шаблоне:
        # 0 — заголовок
        # 1 — "Обозначение (номер) характерной точки / X / Y"
        # 2 — "1 2 3 4"
        # 3+ — данные
        header_rows_count = 3

        # Удаляем только строки данных ниже шапки
        if len(rows) > header_rows_count:
            for i in range(len(rows) - 1, header_rows_count - 1, -1):
                target_table._tbl.remove(rows[i]._tr)

        # Добавляем строки по каждой ЗОУИТ
        for z in zouit_list:
            name = (z.get("name") or "").strip()
            registry_number = (z.get("registry_number") or "").strip()

            if not name and not registry_number:
                continue

            title = name
            if registry_number:
                title = f"{name} ({registry_number})"

            row_cells = target_table.add_row().cells

            # Первая колонка — наименование зоны + реестровый номер
            if len(row_cells) >= 1:
                row_cells[0].text = title

            # Остальные колонки (координаты) оставляем пустыми
            for idx in range(1, len(row_cells)):
                row_cells[idx].text = ""

        logger.info(f"Таблица ЗОУИТ заполнена ({len(zouit_list)} записей)")



    def insert_zouit_blocks(self, doc: Document, zouit_list: List[Dict[str, Any]]) -> None:
        """
        Вставляет текстовые блоки ограничений для ЗОУИТ в раздел 5.
        ОБНОВЛЕННАЯ ВЕРСИЯ с правильными именами файлов + поддержка площадей.
        """
        marker = "{{INSERT_ZOUIT_BLOCKS}}"

        marker_para = None
        for para in doc.paragraphs:
            if marker in para.text:
                marker_para = para
                break

        if marker_para is None:
            logger.warning(f"Маркер {marker!r} не найден для вставки блоков ЗОУИТ")
            return

        # Убираем маркер из текста, но сам параграф оставляем как якорь
        marker_para.text = marker_para.text.replace(marker, "").strip()

        body = marker_para._p.getparent()
        idx = body.index(marker_para._p)

        def add_header_paragraph(name: str, registry_number: str, area: Optional[float]):
            nonlocal body, idx

            p = OxmlElement("w:p")

            # выравнивание по ширине
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "both")
            pPr.append(jc)
            p.append(pPr)

            # run 1: "- "
            r1 = OxmlElement("w:r")
            t1 = OxmlElement("w:t")
            t1.text = "- "
            r1.append(t1)
            p.append(r1)

            # run 2: жирное название
            r2 = OxmlElement("w:r")
            rPr2 = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr2.append(b)
            r2.append(rPr2)

            t2 = OxmlElement("w:t")
            title = name
            if registry_number:
                title += f" ({registry_number})"
            t2.text = title
            r2.append(t2)
            p.append(r2)

            # run 3: площадь
            r3 = OxmlElement("w:r")
            t3 = OxmlElement("w:t")

            if area and area > 0:
                # только целая часть
                try:
                    # математическое округление до целого
                    value = float(area)
                    int_area = int(round(value))
                except Exception:
                    int_area = None

                if int_area is not None:
                    NBSP = "\u00A0"   # неразрывный пробел

                    t3.text = (
                        f". Площадь земельного участка, покрываемая зоной с особыми "
                        f"условиями использования территории, составляет "
                        f"{int_area}{NBSP}кв.м."
                    )
                else:
                    t3.text = ";"
            else:
                t3.text = ";"

            r3.append(t3)
            p.append(r3)

            body.insert(idx + 1, p)
            idx += 1


        for i, z in enumerate(zouit_list, start=1):
            name = z.get("name") or ""
            registry_number = (z.get("registry_number") or "").strip()
            area = z.get("area") or z.get("area_sqm")  # Поддержка обеих полей

            logger.info(f"🔄 Обработка ЗОУИТ {i}/{len(zouit_list)}: {name} ({registry_number}) - площадь: {area}")

            # Заголовок с названием, номером и площадью
            add_header_paragraph(name, registry_number, area)

            # Определяем файл с текстом ограничений (ОБНОВЛЕННАЯ ЛОГИКА)
            block_path = self.get_zouit_block_path(z)
            if block_path is None:
                # Если блока нет — вставляем предупреждение отдельным абзацем
                warn_p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = (
                    f"[ВНИМАНИЕ: Не найден блок ограничений для ЗОУИТ "
                    f"'{name}' ({registry_number})]"
                )
                r.append(t)
                warn_p.append(r)
                body.insert(idx + 1, warn_p)
                idx += 1
                logger.warning(
                    f"❌ Не найден файл блока ограничений для ЗОУИТ "
                    f"{name} ({registry_number})"
                )
            else:
                try:
                    block_doc = Document(str(block_path))
                    # Вставляем содержимое файла сразу после заголовка зоны
                    elements = [deepcopy(el) for el in block_doc.element.body]
                    for el in elements:
                        body.insert(idx + 1, el)
                        idx += 1
                    logger.info(f"✅ Вставлен блок ограничений из {block_path.name}")
                except Exception as ex:
                    logger.error(f"❌ Ошибка загрузки блока {block_path}: {ex}")
                    # Вставляем сообщение об ошибке
                    error_p = OxmlElement("w:p")
                    r = OxmlElement("w:r")
                    t = OxmlElement("w:t")
                    t.text = f"[ОШИБКА: Не удалось загрузить блок {block_path.name}]"
                    r.append(t)
                    error_p.append(r)
                    body.insert(idx + 1, error_p)
                    idx += 1

            # Пустой абзац между зонами
            if i < len(zouit_list):
                empty_p = OxmlElement("w:p")
                body.insert(idx + 1, empty_p)
                idx += 1

        logger.info(f"✅ Вставлено блоков ЗОУИТ: {len(zouit_list)}")


    def insert_coords_table(self, doc: Document, coords: List[Dict[str, Any]]) -> None:
        """Вставляет таблицу координат земельного участка."""
        if not coords:
            logger.info("Координаты отсутствуют, таблица координат не формируется")
            return

        # ВАЖНО: coords уже приходят с правильной нумерацией (как в MID/MIF и интерфейсе)

        # Ищем параграф с маркером [[COORDS_TABLE]]
        p_coords = _find_paragraph_with_text(doc, MARKER_COORDS)
        if not p_coords:
            logger.warning("Маркер [[COORDS_TABLE]] не найден, таблица координат не будет вставлена")
            return

        # Создаём таблицу: сначала 2 строки шапки
        tbl = doc.add_table(rows=2, cols=3)
        try:
            tbl.style = "Table Grid"  # границы таблицы
        except Exception:
            pass

        top = tbl.rows[0].cells
        bot = tbl.rows[1].cells

        # Первая строка шапки
        top[0].text = "Обозначение (номер) характерной точки"
        top[1].text = (
            "Перечень координат характерных точек в системе координат, "
            "используемой для ведения Единого государственного реестра недвижимости"
        )
        top[2].text = ""

        # Вторая строка шапки
        bot[0].text = ""
        bot[1].text = "X"
        bot[2].text = "Y"

        # Объединения ячеек как в образце:
        top[0].merge(bot[0])
        top[1].merge(top[2])

        # Добавляем строки по координатам — строго в порядке из списка
        for coord in coords:
            r = tbl.add_row().cells
            r[0].text = str(coord.get("num") or "").strip()
            r[1].text = _fmt_coord(coord.get("x"))
            r[2].text = _fmt_coord(coord.get("y"))

        # Применяем ширины колонок и выравнивание
        _apply_table_layout(tbl)

        # Вставляем таблицу вместо параграфа с маркером
        _replace_paragraph_with_table(p_coords, tbl)


    def insert_capital_objects_tables(self, doc: Document, capital_objects: List[Dict[str, Any]]) -> None:
        """Вставляет таблицы объектов капитального строительства в раздел 3.1."""
        marker = "[[CAPITAL_OBJECTS_TABLES]]"
        
        marker_para = None
        for para in doc.paragraphs:
            if marker in para.text:
                marker_para = para
                break
        
        if marker_para is None:
            logger.warning(f"Маркер {marker!r} не найден для вставки таблиц ОКС")
            return
        
        # Убираем маркер из текста
        marker_para.text = marker_para.text.replace(marker, "").strip()
        
        body = marker_para._p.getparent()
        idx = body.index(marker_para._p)
        
        if not capital_objects:
            logger.info("ОКС отсутствуют, вставляем 'отсутствует' в таблицу 3.1")

            table = self._create_oks_table(doc, None, 1)

            # --- ВСТАВЛЯЕМ ТЕКСТ В ЯЧЕЙКУ (0,0) ---
            table.rows[0].cells[0].text = "№       Информация отсутствует"

            # --- ОЧИЩАЕМ ТОЛЬКО ТЕ ЯЧЕЙКИ, КОТОРЫЕ ДОЛЖНЫ БЫТЬ ПУСТЫМИ ---
            # (0,1) → пустая
            table.rows[0].cells[1].text = ""

            # (1,1) — НЕ трогаем! текст должен остаться из шаблона

            # (2,1) → пустая
            table.rows[2].cells[1].text = ""

            body.insert(idx + 1, table._tbl)
            return
        else:
            # Если объекты есть — создаём таблицу для каждого
            logger.info(f"Создаём {len(capital_objects)} таблиц ОКС")
            for i, obj in enumerate(capital_objects, start=1):
                table = self._create_oks_table(doc, obj, i)
                body.insert(idx + 1, table._tbl)
                idx += 1
                
                # Пустой абзац между таблицами (кроме последней)
                if i < len(capital_objects):
                    empty_p = OxmlElement("w:p")
                    body.insert(idx + 1, empty_p)
                    idx += 1
        
        logger.info("Таблицы ОКС вставлены в раздел 3.1")


    def _create_oks_table(self, doc: Document, obj: Optional[Dict[str, Any]], num: int) -> Table:
        """Создаёт одну таблицу объекта капитального строительства."""
        # Создаём таблицу 3×2 БЕЗ СТИЛЯ (без границ)
        table = doc.add_table(rows=3, cols=2)
        
        # Ширина колонок: 8.5 см каждая
        col_width = Cm(8.5)
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
        
        # Функция добавления тонкой нижней границы
        def add_bottom_border(cell):
            """Добавляет тонкую нижнюю границу к ячейке"""
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '2')  # Очень тонкая граница
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
        
        # === СТРОКА 1 === #
        cell_0_0 = table.rows[0].cells[0]
        cell_0_1 = table.rows[0].cells[1]
        
        # Выравнивание
        cell_0_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_0_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # [0,0]: "№X" или пустая
        p1 = cell_0_0.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if obj:
            # Если есть объект — выводим номер
            run1 = p1.add_run(f"№{num}")
            run1.font.name = "Times New Roman"
            run1.font.size = Pt(12)
        # Иначе ячейка остаётся пустой
        
        # Тонкая граница снизу
        add_bottom_border(cell_0_0)
        
        # [0,1]: "НЕОБХОДИМО ЗАПОЛНИТЬ ДАННЫЕ" или пустая
        p2 = cell_0_1.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if obj:
            # Если есть объект — красный текст
            run2 = p2.add_run("НЕОБХОДИМО ЗАПОЛНИТЬ ДАННЫЕ")
            run2.font.name = "Times New Roman"
            run2.font.size = Pt(12)
            run2.font.color.rgb = RGBColor(255, 0, 0)
            run2.bold = True
        # Иначе ячейка остаётся пустой
        
        # Тонкая граница снизу
        add_bottom_border(cell_0_1)
        
        # === СТРОКА 2 === #
        cell_1_0 = table.rows[1].cells[0]
        cell_1_1 = table.rows[1].cells[1]
        
        # Выравнивание
        cell_1_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_1_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # [1,0]: "(согласно чертежу(ам)...)" - БЕЗ ГРАНИЦ
        p3 = cell_1_0.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("(согласно чертежу(ам) планировки территории)")
        run3.font.name = "Times New Roman"
        run3.font.size = Pt(8)
        run3.font.color.rgb = RGBColor(128, 128, 128)
        
        # [1,1]: "(назначение объекта...)" - БЕЗ ГРАНИЦ
        p4 = cell_1_1.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run("(назначение объекта капитального строительства, этажность, высотность, общая площадь,  площадь застройки)")
        run4.font.name = "Times New Roman"
        run4.font.size = Pt(8)
        run4.font.color.rgb = RGBColor(128, 128, 128)
        
        # === СТРОКА 3 === #
        cell_2_0 = table.rows[2].cells[0]
        cell_2_1 = table.rows[2].cells[1]
        
        # Выравнивание
        cell_2_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_2_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # [2,0]: "инвентаризационный или кадастровый номер" - БЕЗ ГРАНИЦ
        p5 = cell_2_0.paragraphs[0]
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = p5.add_run("инвентаризационный или кадастровый номер")
        run5.font.name = "Times New Roman"
        run5.font.size = Pt(12)
        
        # [2,1]: Кадастровый номер или пустая
        p6 = cell_2_1.paragraphs[0]
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if obj and obj.get("cadnum"):
            # Если есть кадастровый номер — выводим его
            run6 = p6.add_run(obj["cadnum"])
            run6.font.name = "Times New Roman"
            run6.font.size = Pt(12)
        # Иначе ячейка остаётся пустой
        
        # Тонкая граница снизу
        add_bottom_border(cell_2_1)
        
        return table


    def generate(self, gp_data: Dict[str, Any], output_path: str) -> str:
        """Генерирует градостроительный план."""
        logger.info("🚀 Начало генерации градплана")

        # --- 1. Рендер шаблона через docxtpl ---
        tpl = DocxTemplate(self.template_path)
        context = self.prepare_context(gp_data)
        tpl.render(context)

        temp_path = str(Path(output_path).with_suffix(".tmp.docx"))
        tpl.save(temp_path)

        # Загружаем как обычный Document для низкоуровневых операций
        doc = Document(temp_path)

        # --- 2. Таблица координат ---
        parcel = gp_data.get("parcel") or {}
        coords = parcel.get("coordinates") or []
        if coords:
            self.insert_coords_table(doc, coords)
        else:
            logger.info("Координаты участка в данных отсутствуют")

        # --- 3. Таблицы ОКС ---
        capital_objects = gp_data.get("capital_objects") or []
        self.insert_capital_objects_tables(doc, capital_objects)

        # --- 4. Блоки территориальных зон ---
        zone = gp_data.get("zone") or {}
        zone_code = zone.get("code")
        if zone_code:
            # ВРИ
            vri_block = self.load_zone_block(zone_code, "vri")
            if vri_block:
                self.insert_block_at_marker(doc, "{{INSERT_ZONE_VRI}}", vri_block)
            else:
                logger.warning(f"Не найден блок ВРИ для зоны {zone_code}")

            # Параметры
            params_block = self.load_zone_block(zone_code, "params")
            if params_block:
                self.insert_block_at_marker(doc, "{{INSERT_ZONE_PARAMS}}", params_block)
            else:
                logger.warning(f"Не найден блок параметров для зоны {zone_code}")

        # --- 4б. Блок АГО (п.12) ---
        ago_index = gp_data.get("ago_index")
        self.insert_ago_block(doc, ago_index)

        # --- 5. ЗОУИТ (ОБНОВЛЕННАЯ ЛОГИКА С ПЛОЩАДЯМИ) ---
        zouit_list = gp_data.get("zouit") or []
        if zouit_list:
            logger.info(f"📋 Обработка {len(zouit_list)} ЗОУИТ")
            self.fill_zouit_table(doc, zouit_list)
            self.insert_zouit_blocks(doc, zouit_list)
        else:
            logger.info("ЗОУИТ для участка отсутствуют")

        # --- 6. Сохранение результата ---
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))

        # Удаляем временный файл
        try:
            os.remove(temp_path)
        except OSError:
            pass

        logger.info(f"✅ ГПЗУ успешно сформирован: {out_path}")
        return str(out_path)


# ================ ФУНКЦИЯ ДЛЯ ТЕСТИРОВАНИЯ ================ #

def test_zouit_mapping():
    """
    Функция для тестирования сопоставления ЗОУИТ с реальными файлами
    """
    
    test_cases = [
        "Водоохранная зона",
        "Санитарно-защитная зона", 
        "Охранная зона объектов электросетевого хозяйства",
        "Охранная зона ВЛ-35 кВ",
        "Газораспределительная сеть",
        "Тепловые сети", 
        "Линии связи",
        "Объекты культурного наследия",
        "Зона затопления",
        "Пункты наблюдения окружающей среды",
        "Геодезическая сеть",
        "Магистральный трубопровод",
        "Радиотехнический объект",
        "Приаэродромная территория",
        "четвертая подзона приаэродромной территории",
        "Санитарная зона источников водоснабжения",
        "Железная дорога"
    ]
    
    # Создаем тестовый экземпляр
    try:
        base_dir = Path(__file__).resolve().parent.parent
        template_path = base_dir / "templates" / "gpzu_template.docx"
        
        builder = GPBuilder(str(template_path))
        results = builder.test_zouit_matching(test_cases)
        
        print("\n=== РЕЗУЛЬТАТЫ ТЕСТИРОВАНИЯ СОПОСТАВЛЕНИЯ ЗОУИТ ===")
        print(f"{'НАЗВАНИЕ ЗОУИТ':<60} {'НАЙДЕННЫЙ ФАЙЛ'}")
        print("=" * 100)
        
        for name, filename in results.items():
            status = "✅" if filename != "НЕ НАЙДЕН" else "❌"
            print(f"{status} {name:<58} -> {filename}")
            
        return results
        
    except Exception as e:
        print(f"Ошибка тестирования: {e}")
        return {}


def generate_gp_document(gp_data: Dict[str, Any], output_path: str) -> str:
    """
    Утилитная функция для генерации ГПЗУ "в одно действие".
    ОБНОВЛЕННАЯ ВЕРСИЯ с правильными именами файлов ЗОУИТ.
    """
    base_dir = Path(__file__).resolve().parent.parent
    template_path = base_dir / "templates" / "gpzu_template.docx"
    builder = GPBuilder(str(template_path))
    return builder.generate(gp_data, output_path)


if __name__ == "__main__":
    # Запуск тестирования
    test_zouit_mapping()
