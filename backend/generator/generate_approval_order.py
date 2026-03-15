# backend/generator/generate_approval_order.py
"""
Генератор приказа об утверждении шаблона решения РРР.
Создаёт DOCX-документ с описанием всех переменных, условных блоков,
типов объектов, реквизитов и листом согласования.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

CONFIG_DIR = Path(__file__).resolve().parent.parent / "config"

# ---------------------------------------------------------------------------
# Описания переменных шаблона
# ---------------------------------------------------------------------------

VARIABLE_DESCRIPTIONS = [
    # Решение
    ("OUT_NUMBER",              "Номер решения",                           "Карточка РРР → поле «Номер решения»",              "123/2026"),
    ("OUT_DATE",                "Дата решения (ДД.ММ.ГГГГ)",               "Карточка РРР → поле «Дата решения»",               "15.03.2026"),
    # Заявитель
    ("APPLICANT",               "Наименование/ФИО заявителя",              "Организация (ЮЛ) или ФИО физлица (ФЛ)",            "ООО «Ромашка»"),
    ("APPLICANT_DETAILS",       "Реквизиты заявителя (ИНН, ОГРН, адрес)", "Формируется из полей карточки",                    "ИНН 1234567890, ОГРН 1021234567890, г. Новокузнецк"),
    # Объект
    ("OBJECT_TYPE_FULL",        "Полное наименование вида объекта",        "Из object_types.json по номеру пункта",            "Линии электропередачи классом напряжения до 35 кВ..."),
    ("OBJECT_NAME",             "Конкретное наименование объекта",         "Карточка РРР → поле «Наименование объекта»",       "ВЛ-10 кВ фидер Северный"),
    ("AREA",                    "Площадь объекта (форматированная)",       "Карточка РРР → поле «Площадь», формат 1 024,46",   "1 024,46"),
    ("AREA_NUM",                "Площадь объекта (число)",                 "Карточка РРР → поле «Площадь»",                    "1024.46"),
    ("LOCATION",                "Местоположение объекта",                  "Карточка РРР → поле «Местоположение»",             "г. Новокузнецк, ул. Ленина"),
    # Срок
    ("TERM_TEXT",               "Срок прописью",                           "Вычисляется из TERM_MONTHS",                       "36 (тридцать шесть) месяцев"),
    ("END_DATE",                "Дата окончания (ДД.ММ.ГГГГ)",             "decision_date + term_months - 1 день",             "14.03.2029"),
    # Флаги (условные блоки)
    ("is_linear",               "Флаг: линейный объект",                   "Из object_types.json (is_linear)",                 "True / False"),
    ("is_nto",                  "Флаг: нестационарный объект (НТО)",       "Из object_types.json (is_nto)",                    "True / False"),
    ("has_payment",             "Флаг: объект платный",                    "Из конфига или ручное переопределение",            "True / False"),
    ("has_red_lines",           "Флаг: красные линии пересекают объект",   "Из результатов пространственного анализа",         "True / False"),
    ("proezd_agreement",        "Согласование примыкания (п.12)",          "Карточка РРР → поле «Согласование проезда»",       "ГИБДД ГУ МВД РФ"),
    ("payment_formula",         "Формула расчёта платы",                   "standard или lep (для п.5 ЛЭП)",                   "standard"),
    # Оплата
    ("N_PAYMENT",               "Номер пункта об оплате",                  "Динамическая нумерация (3, если есть оплата)",     "3"),
    ("is_attrac",               "Флаг: объект — аттракцион (п.24)",        "Из object_types.json (is_attrac)",                 "True / False"),
    ("N_ATTRAC",                "Номер пункта об аттракционе",             "Динамическая нумерация; 0 если не аттракцион",     "4 или 0"),
    ("PAYMENT_YEARLY",          "Годовая плата (форматированная)",         "Вычисляется из площади и формулы",                 "13 649,70"),
    ("PAYMENT_TOTAL",           "Итоговая плата за период (форматир.)",    "PAYMENT_YEARLY × (дней / 365)",                    "40 949,10"),
    ("PAYMENT_TOTAL_WORDS",     "Итоговая плата прописью",                 "Вычисляется из PAYMENT_TOTAL",                     "сорок тысяч девятьсот сорок девять рублей 10 копеек"),
    ("PERIOD_DAYS",             "Количество дней действия разрешения",     "end_date − decision_date",                         "1094"),
    ("KI_FORMULA",              "Формула КИ (произведение коэффициентов)", "Из payment_config.json (lep.ki_formula)",          "1,037 × 1,040 × 1,055 × 1,045 × 1,045 × 1,04"),
    # Реквизиты
    ("REQ_BANK",                "Банк получателя",                         "Из payment_config.json (requisites.bank)",         "ОКЦ №5 СибГУ..."),
    # Нумерация пунктов (динамическая)
    ("N_TERMINATION",           "Номер пункта «Прекращение действия»",     "Динамическая нумерация (зависит от has_payment)",  "4 или 3"),
    ("N_LIQUIDATION",           "Номер пункта «В случае ликвидации»",      "Динамическая нумерация",                           "5"),
    ("N_EARTHWORKS",            "Номер пункта «Земляные работы»",          "Динамическая нумерация",                           "6"),
    ("N_GEODESY",               "Номер пункта «Геодезическая съёмка»",     "Динамическая нумерация",                           "7"),
    ("N_RSO",                   "Номер пункта «Согласования РСО»",         "Динамическая нумерация",                           "8"),
    ("N_THIRD_PARTIES",         "Номер пункта «Права третьих лиц»",        "Динамическая нумерация",                           "9"),
    ("N_PREV_DECISIONS",        "Номер пункта «Ранее выданные РРР»",       "0 если нет ранее выданных решений",                "10 или 0"),
    ("N_VEGETATION",            "Номер пункта «Вырубка насаждений»",       "Динамическая нумерация",                           "11"),
    ("N_CLEANUP",               "Номер пункта «Очистка от деревьев»",      "Динамическая нумерация",                           "12"),
    ("N_LINER",                 "Номер пункта «Линейные объекты»",         "0 если не линейный объект",                        "13 или 0"),
    ("N_END_TERM",              "Номер пункта «По окончании срока»",       "Динамическая нумерация",                           "14"),
    ("N_MAINTENANCE",           "Номер пункта «Содержание территории»",    "Динамическая нумерация",                           "15"),
    ("N_PAVEMENT",              "Номер пункта «Восстановление покрытия»",  "Динамическая нумерация",                           "16"),
    ("N_CONTROL",               "Номер пункта «Контроль»",                 "Динамическая нумерация",                           "17"),
    ("N_TERM_PAYMENT",          "Номер подпункта «Оплата» в п.Прекращение","5 если платный, 0 если бесплатный",               "5 или 0"),
    ("N_TERM_NTO",              "Номер подпункта «аттракцион» в п.Прекращение (управляется is_attrac, не is_nto)", "Следующий после N_TERM_PAYMENT", "6 или 5"),
    # Пространственный анализ
    ("QUARTERS",                "Кадастровые кварталы",                    "Результат пространственного анализа",              "42:17:0101001"),
    # Ранее выданные решения
    ("PREV_DECISIONS",          "Список ранее выданных решений РРР",       "Карточка РРР → вкладка «Ранее выданные»",          "[{decision_number, decision_date, end_date, ...}]"),
    ("CURRENT_USER",            "Текущий пользователь (исполнитель)",      "Из сессии авторизации",                            "Иванова А.П."),
]

# ---------------------------------------------------------------------------
# Условные блоки шаблона
# ---------------------------------------------------------------------------

CONDITIONAL_BLOCKS = [
    (
        "{% if is_linear %}",
        "Линейный объект (п.1,2,3,5,6,7,8,32)",
        "статьи 51 Градостроительного кодекса Российской Федерации, постановления Правительства Российской Федерации от 12.11.2020 №1816 «Об утверждении перечня случаев, при которых для строительства, реконструкции линейного объекта не требуется подготовка документации по планировке территории»,"
    ),
    (
        "{% if is_nto %}",
        "Нестационарный объект (п.19,23,24,25,36)",
        ", постановления администрации города Новокузнецка от 11.04.2018 № 63 «Об утверждении Положения об организации размещения нестационарных торговых объектов на территории Новокузнецкого городского округа»"
    ),
    (
        "{% if proezd_agreement %}",
        "Заполнено поле «Согласование примыкания»",
        "{{ proezd_agreement }} — вставляется произвольный текст согласования примыкания (вводится вручную в карточке РРР)"
    ),
    (
        "{% if APPLICANT_DETAILS %}",
        "Есть реквизиты заявителя (ИНН/ОГРН/адрес/паспорт)",
        "({{ APPLICANT_DETAILS }}) — вставляется в п.1 после наименования заявителя"
    ),
    (
        "{% if has_payment %}",
        "Объект платный (has_payment == True)",
        "{{ N_PAYMENT }}.\tПлата за размещение Объекта вносится заявителем в {{ REQ_BANK }} БИК 013207212 на счет 40102810745370000032. Получатель платежа: УФК по Кемеровской области – Кузбассу (Комитет градостроительства и архитектуры администрации города Новокузнецка) ИНН 4217121181 КПП 421701001, счет 03100643000000013900; ОКТМО 32731000. Назначение и код платежа: Прочие неналоговые доходы бюджетов городских округов (плата за разрешение на размещение объекта), КБК 906 117 05040 04 0040 180, в размере согласно приложению №2 к настоящему Решению.\n(следующий абзац) Плата за размещение Объекта вносится единовременно за весь период размещения объекта в течение 20 рабочих дней со дня направления Комитетом градостроительства и архитектуры администрации города Новокузнецка (далее – Комитет) настоящего Решения."
    ),
    (
        "{%p if is_attrac %}",
        "Объект — аттракцион (п.24)",
        "{{ N_ATTRAC }}. Лицо, получившее решение о разрешении размещения указанного объекта, приступает к использованию земельного участка только после получения свидетельства о государственной регистрации аттракциона, а также акта оценки его технического состояния (технического освидетельствования), подтверждающего соответствие такого аттракциона перечню требований к техническому состоянию и эксплуатации аттракционов, со дня выдачи которого прошло не более 12 месяцев (для сезонных аттракционов, изготовленных и введённых в эксплуатацию до вступления в силу технического регламента Евразийского экономического союза «О безопасности аттракционов»)."
    ),
    (
        "{%p if has_payment %}\nв п. «Прекращение»",
        "Объект платный (подпункт в блоке Прекращения)",
        "{{ N_TERM_PAYMENT }})\tневнесения платы за размещение объекта в срок, установленный пунктом {{ N_PAYMENT }} настоящего решения."
    ),
    (
        "{% if is_attrac %}\nв п. «Прекращение»",
        "Объект — аттракцион (подпункт в блоке Прекращения). ⚠ Переменная N_TERM_NTO, управляется is_attrac",
        "{{ N_TERM_NTO }})\tпоступления в уполномоченный орган информации об использовании земель или земельных участков для размещения сезонного аттракциона при отсутствии свидетельства о государственной регистрации аттракциона, а также акта оценки его технического состояния (технического освидетельствования), подтверждающего соответствие такого аттракциона перечню требований к техническому состоянию и эксплуатации аттракционов, со дня выдачи которого прошло не более 12 месяцев (для сезонных аттракционов, изготовленных и введенных в эксплуатацию до вступления в силу технического регламента Евразийского экономического союза «О безопасности аттракционов»)"
    ),
    (
        "{% if has_red_lines %}",
        "Красные линии пересекают объект (пространственный анализ)",
        "При размещении объекта необходимо учитывать документацию по проекту планировки улично-дорожной сети Новокузнецкого городского округа, утвержденную распоряжением администрации города Новокузнецка от 19.10.2016 №1766"
    ),
    (
        "{%p for d in PREV_DECISIONS %}",
        "Есть ранее выданные решения РРР (список непустой)",
        "{{ N_PREV_DECISIONS }}. Размещение Объекта не должно препятствовать размещению объекта вида «{{ d.object_type }}», указанному в ранее выданном решении Комитета от {{ d.decision_date }} №Н{{ d.decision_number }} «О разрешении размещения объекта без предоставления земельного участка и установления сервитута, публичного сервитута» со сроком действия до {{ d.end_date }}. Получатель вышеуказанного решения — {{ d.applicant }}."
    ),
    (
        "{% if is_linear %}\n(пункт N_LINER)",
        "Линейный объект — отдельный пункт об техприсоединении",
        "{{ N_LINER }}.\tРазмещаемый Объект предназначен для подключения (технологического присоединения) объектов капитального строительства к сетям инженерно-технического обеспечения, предоставление земельных участков, на которых или под поверхностью которых размещен указанный объект, не влечет за собой принудительные снос или демонтаж указанного Объекта."
    ),
    (
        "{%p if has_payment %}\nПриложение №2",
        "Объект платный — весь лист «Расчёт платы»",
        "Приложение №2 к решению: РАСЧЕТ ПЛАТЫ ЗА РАЗМЕЩЕНИЕ ОБЪЕКТА — таблица с формулой, расшифровкой величин, расчётом за период и итоговой суммой прописью"
    ),
    (
        "{% if payment_formula == 'lep' %}\nв таблице Прил.2",
        "Формула ЛЭП (объект п.5 — ЛЭП до 35 кВ)",
        "Формула: П = П2020 × КИ × S; расшифровка: П2020 – годовой размер платы за 1 кв.м., КИ – произведение коэффициентов инфляции КИ = КИ2021 × КИ2022 × КИ2023 × КИ2024 × КИ2025 × КИ2026"
    ),
    (
        "{% else %} (payment_formula != lep)\nв таблице Прил.2",
        "Стандартная формула (все объекты кроме п.5)",
        "Формула: П = Су × Нст × S; расшифровка: Су – удельный показатель кадастровой стоимости земли, Нст – ставка платы (1,5%)"
    ),
]

# ---------------------------------------------------------------------------
# Подписанты листа согласования
# ---------------------------------------------------------------------------

APPROVERS = [
    "Начальник правового отдела",
    "Начальник отдела архитектуры",
    "Заместитель начальника отдела",
    "Начальник отдела ИСОГД",
    "Начальник экономико-аналитического отдела",
    "Заместитель председателя",
    "Главный архитектор",
    "Председатель Комитета",
]

# ---------------------------------------------------------------------------
# Вспомогательные функции форматирования
# ---------------------------------------------------------------------------


def _set_font(run, size_pt: int = 12, bold: bool = False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Times New Roman"


def _para(doc: Document, text: str = "", align=WD_ALIGN_PARAGRAPH.LEFT,
          size_pt: int = 12, bold: bool = False) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        _set_font(run, size_pt, bold)
    return p


def _heading(doc: Document, text: str, level: int = 1):
    p = _para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER,
              size_pt=13 if level == 1 else 12, bold=True)
    return p


def _set_table_borders(table):
    """Добавляет видимые границы ко всем ячейкам таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_text(cell, text: str, size_pt: int = 10, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    _set_font(run, size_pt, bold)


def _set_col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ---------------------------------------------------------------------------
# Построение документа
# ---------------------------------------------------------------------------


def generate_approval_order(output_path: str) -> str:
    """
    Сгенерировать приказ об утверждении шаблона решения РРР.

    Args:
        output_path: Путь для сохранения DOCX

    Returns:
        Путь к созданному файлу
    """
    doc = Document()

    # Настройка полей страницы
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # -------------------------------------------------------------------
    # РАЗДЕЛ 1: Шапка приказа
    # -------------------------------------------------------------------
    _para(doc, "КОМИТЕТ ГРАДОСТРОИТЕЛЬСТВА И АРХИТЕКТУРЫ",
          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13, bold=True)
    _para(doc, "АДМИНИСТРАЦИИ ГОРОДА НОВОКУЗНЕЦКА",
          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13, bold=True)

    _para(doc)  # пустая строка

    _para(doc, "ПРИКАЗ № ___ от «___» __________ 2026 г.",
          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13, bold=True)

    _para(doc)

    _para(doc,
          "Об утверждении шаблона решения о разрешении размещения объектов\n"
          "без предоставления земельного участка (РРР)",
          align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, bold=True)

    _para(doc)

    body_text = (
        "В целях автоматизированного формирования решений о разрешении размещения объектов "
        "без предоставления земельного участка на территории города Новокузнецка ПРИКАЗЫВАЮ:"
    )
    p_body = _para(doc, body_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size_pt=12)

    _para(doc)

    for num, clause in [
        ("1.", "Утвердить шаблон решения о разрешении размещения объектов согласно Приложению №1."),
        ("2.", "Настоящий приказ вступает в силу с момента подписания."),
    ]:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.clear()
        run = p.add_run(f"{num} {clause}")
        _set_font(run, 12)

    _para(doc)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sign = p_sign.add_run("Председатель Комитета")
    _set_font(run_sign, 12)
    p_sign.add_run("\t\t\t_______________ / В.С. Левченко /")

    doc.add_page_break()

    # -------------------------------------------------------------------
    # РАЗДЕЛ 2: Приложение №1 — Описание шаблона
    # -------------------------------------------------------------------
    _heading(doc, "ПРИЛОЖЕНИЕ №1 к Приказу № ___ от «___» __________ 2026 г.")
    _heading(doc, "Описание шаблона решения о разрешении размещения объектов")

    _para(doc)

    # --- Таблица 1: Переменные шаблона ---
    _para(doc, "Таблица 1. Переменные шаблона", size_pt=12, bold=True)

    headers_t1 = ["Переменная", "Описание", "Источник", "Пример"]
    col_widths_t1 = [4.5, 5.5, 5.5, 3.5]

    t1 = doc.add_table(rows=1 + len(VARIABLE_DESCRIPTIONS), cols=4)
    _set_table_borders(t1)

    # Заголовок
    for i, h in enumerate(headers_t1):
        _cell_text(t1.rows[0].cells[i], h, size_pt=10, bold=True)
        t1.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные
    for row_idx, (var, desc, source, example) in enumerate(VARIABLE_DESCRIPTIONS, start=1):
        row = t1.rows[row_idx]
        _cell_text(row.cells[0], var, size_pt=9)
        _cell_text(row.cells[1], desc, size_pt=9)
        _cell_text(row.cells[2], source, size_pt=9)
        _cell_text(row.cells[3], example, size_pt=9)

    _set_col_widths(t1, col_widths_t1)

    _para(doc)

    # --- Таблица 2: Условные блоки ---
    _para(doc, "Таблица 2. Условные блоки шаблона", size_pt=12, bold=True)

    headers_t2 = ["Условие", "Когда активно", "Что показывается"]
    col_widths_t2 = [4.0, 5.0, 10.0]

    t2 = doc.add_table(rows=1 + len(CONDITIONAL_BLOCKS), cols=3)
    _set_table_borders(t2)

    for i, h in enumerate(headers_t2):
        _cell_text(t2.rows[0].cells[i], h, size_pt=10, bold=True)
        t2.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (cond, when, shows) in enumerate(CONDITIONAL_BLOCKS, start=1):
        row = t2.rows[row_idx]
        _cell_text(row.cells[0], cond, size_pt=9)
        _cell_text(row.cells[1], when, size_pt=9)
        _cell_text(row.cells[2], shows, size_pt=9)

    _set_col_widths(t2, col_widths_t2)

    _para(doc)

    # --- Таблица 3: Типы объектов ---
    _para(doc, "Таблица 3. Типы объектов (Постановление Правительства РФ №1300)", size_pt=12, bold=True)

    with open(CONFIG_DIR / "object_types.json", encoding="utf-8") as f:
        object_types = json.load(f)

    headers_t3 = ["№ п.", "Наименование", "Линейный", "НТО", "Платный по умолч.", "Макс. срок"]
    col_widths_t3 = [1.2, 8.0, 2.0, 1.5, 2.5, 2.3]

    t3 = doc.add_table(rows=1 + len(object_types), cols=6)
    _set_table_borders(t3)

    for i, h in enumerate(headers_t3):
        _cell_text(t3.rows[0].cells[i], h, size_pt=9, bold=True)
        t3.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, obj in enumerate(object_types, start=1):
        row = t3.rows[row_idx]
        _cell_text(row.cells[0], f"п.{obj['number']}", size_pt=9)
        _cell_text(row.cells[1], obj["short_name"], size_pt=9)
        _cell_text(row.cells[2], "Да" if obj.get("is_linear") else "Нет", size_pt=9)
        _cell_text(row.cells[3], "Да" if obj.get("is_nto") else "Нет", size_pt=9)
        by_name = obj.get("payment_by_name") or obj.get("payment_by_name_oil") or obj.get("payment_by_name_cafe")
        if obj.get("has_payment_default"):
            payment_label = "Да"
        elif by_name:
            payment_label = "Нет*"
        else:
            payment_label = "Нет"
        _cell_text(row.cells[4], payment_label, size_pt=9)
        _cell_text(row.cells[5], f"{obj.get('max_term_years', 3)} лет", size_pt=9)

    _set_col_widths(t3, col_widths_t3)

    p_note = _para(doc,
                   "* Нет* — платность определяется по наименованию объекта "
                   "(п.4 — парковки/стоянки; п.6 — нефтепроводы; п.19 — кафе/общественное питание).",
                   size_pt=9)

    _para(doc)

    # --- Таблица 4: Реквизиты ---
    _para(doc, "Таблица 4. Реквизиты для платежей", size_pt=12, bold=True)

    with open(CONFIG_DIR / "payment_config.json", encoding="utf-8") as f:
        payment_config = json.load(f)

    req = payment_config.get("requisites", {})
    req_rows = [
        ("Получатель",              req.get("recipient", "")),
        ("ИНН получателя",          req.get("inn", "")),
        ("КПП получателя",          req.get("kpp", "")),
        ("Банк",                    req.get("bank", "")),
        ("БИК",                     req.get("bik", "")),
        ("Расчётный счёт",          req.get("account", "")),
        ("Корреспондентский счёт",  req.get("corr_account", "")),
        ("ОКТМО",                   req.get("oktmo", "")),
        ("КБК",                     req.get("kbk", "")),
        ("Назначение платежа",      req.get("purpose", "")),
    ]

    t4 = doc.add_table(rows=1 + len(req_rows), cols=2)
    _set_table_borders(t4)

    for i, h in enumerate(["Поле", "Значение"]):
        _cell_text(t4.rows[0].cells[i], h, size_pt=10, bold=True)
        t4.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (field, value) in enumerate(req_rows, start=1):
        row = t4.rows[row_idx]
        _cell_text(row.cells[0], field, size_pt=10)
        _cell_text(row.cells[1], value, size_pt=10)

    _set_col_widths(t4, [5.0, 14.0])

    _para(doc)

    # --- Таблица 5: Формулы расчёта ---
    _para(doc, "Таблица 5. Формулы расчёта платы", size_pt=12, bold=True)

    std = payment_config.get("standard", {})
    lep = payment_config.get("lep", {})

    formula_rows = [
        ("Стандартная формула",
         f"П_год = Су × Нст × S; П_итого = П_год × (дней / 365)\n"
         f"где Су = {std.get('su', '')} руб./кв.м ({std.get('su_description', '')}),\n"
         f"Нст = {std.get('nst', '')} ({std.get('nst_description', '')})",
         "Все объекты кроме п.5 (ЛЭП)"),
        ("Формула ЛЭП (п.5)",
         f"П_год = {lep.get('base_rate', '')} × КИ × S; П_итого = П_год × (дней / 365)\n"
         f"где КИ = {lep.get('ki', '')} = {lep.get('ki_formula', '')}\n"
         f"({lep.get('ki_description', '')})",
         "Только п.5 — Линии электропередачи до 35 кВ"),
    ]

    t5 = doc.add_table(rows=1 + len(formula_rows), cols=3)
    _set_table_borders(t5)

    for i, h in enumerate(["Формула", "Описание", "Применяется для"]):
        _cell_text(t5.rows[0].cells[i], h, size_pt=10, bold=True)
        t5.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (name, formula, applies_to) in enumerate(formula_rows, start=1):
        row = t5.rows[row_idx]
        _cell_text(row.cells[0], name, size_pt=10, bold=True)
        _cell_text(row.cells[1], formula, size_pt=9)
        _cell_text(row.cells[2], applies_to, size_pt=9)

    _set_col_widths(t5, [4.0, 10.0, 5.0])

    doc.add_page_break()

    # -------------------------------------------------------------------
    # РАЗДЕЛ 3: Лист согласования
    # -------------------------------------------------------------------
    _heading(doc, "ЛИСТ СОГЛАСОВАНИЯ")

    _para(doc)

    _para(doc, "Наименование: Шаблон решения о разрешении размещения объектов "
               "без предоставления земельного участка (РРР)", size_pt=12)

    _para(doc)

    t_sign = doc.add_table(rows=1 + len(APPROVERS), cols=4)
    _set_table_borders(t_sign)

    for i, h in enumerate(["Должность", "ФИО", "Подпись", "Дата"]):
        _cell_text(t_sign.rows[0].cells[i], h, size_pt=11, bold=True)
        t_sign.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, approver in enumerate(APPROVERS, start=1):
        row = t_sign.rows[row_idx]
        _cell_text(row.cells[0], approver, size_pt=11)
        # Остальные ячейки оставляем пустыми для подписей
        for cell in row.cells[1:]:
            cell.text = ""
            # Минимальная высота строки для подписи
            tc = cell._tc
            trPr = tc.getparent().get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "800")
            trPr.append(trHeight)

    _set_col_widths(t_sign, [8.0, 5.0, 3.0, 3.0])

    _para(doc)
    _para(doc, "Документ сформирован автоматически системой ГПЗУ-Web.", size_pt=9)

    # Сохранение
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return str(out)


# ---------------------------------------------------------------------------
# Точка входа (запуск напрямую)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "uploads/approval_order.docx"
    result = generate_approval_order(output)
    print(f"Документ сохранён: {result}")
