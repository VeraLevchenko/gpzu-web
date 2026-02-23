# backend/parsers/rrr_application_parser.py
"""
Парсер заявления DOCX для модуля РРР.
Извлекает данные заявителя из таблиц документа.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from io import BytesIO
from typing import Optional, Tuple
import re

from docx import Document


@dataclass
class RRRApplicationData:
    """Результат разбора заявления РРР."""
    # Юрлицо
    org_name: Optional[str] = None
    inn: Optional[str] = None
    ogrn: Optional[str] = None
    org_address: Optional[str] = None
    # Физлицо
    person_name: Optional[str] = None
    person_passport: Optional[str] = None
    person_address: Optional[str] = None
    # Реквизиты заявления
    app_number: Optional[str] = None
    app_date: Optional[date] = None
    app_date_text: Optional[str] = None
    # Срок
    term_months: Optional[int] = None


def _load_doc(doc_bytes: bytes) -> Document:
    return Document(BytesIO(doc_bytes))


def _extract_from_tables(doc: Document) -> dict:
    """
    Извлекает все доступные поля из таблиц документа.
    """
    result = {}
    months = {
        "января": 1, "февраля": 2, "марта": 3, "апреля": 4, "мая": 5,
        "июня": 6, "июля": 7, "августа": 8, "сентября": 9, "октября": 10,
        "ноября": 11, "декабря": 12,
    }

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if not cells:
                continue

            row_text = " ".join(cells).lower()

            # Номер заявления
            for c in cells:
                m = re.search(r"[№N]\s*:?\s*([0-9]{3,})", c)
                if m and "app_number" not in result:
                    result["app_number"] = m.group(1)

            # Дата заявления
            for c in cells:
                if "«" in c and "»" in c and "г" in c and "app_date" not in result:
                    result["app_date_text"] = c.strip()
                    try:
                        day_part = c.split("«", 1)[1].split("»", 1)[0]
                        day = int(day_part.strip().strip(". "))
                        rest = c.split("»", 1)[1].replace("г.", "").replace("г", "").strip()
                        parts = rest.split()
                        if len(parts) >= 2:
                            month = months.get(parts[0].lower())
                            year = int(parts[1])
                            if day and month and year:
                                result["app_date"] = date(year, month, day)
                    except Exception:
                        pass

            # Юрлицо: полное наименование (1.2.1)
            if cells[0].startswith("1.2.1"):
                if len(cells) >= 3 and cells[2]:
                    result["org_name"] = cells[2].strip()

            # ФИО физлица (1.1.1) — нумерованный формат
            if cells[0].startswith("1.1.1"):
                if len(cells) >= 3 and cells[2]:
                    result["person_name"] = cells[2].strip()

            # ФИО физлица — формат ЕПГУ «ФИО: …;»
            if "person_name" not in result and "фио" in row_text:
                for c in cells:
                    m = re.search(r'ФИО\s*:\s*([^;]+)', c)
                    if m:
                        name = m.group(1).strip().rstrip(';').strip()
                        if name and len(name) > 3:
                            result["person_name"] = name
                            break

            # Полное наименование из свободного текста ячейки
            # Формат: "Полное наименование: ООО "НАЗВАНИЕ"; ИНН: ..."
            if "org_name" not in result:
                for c in cells:
                    m = re.search(
                        r'[Пп]олное\s+наименование\s*:\s*(.+?)(?:\s*;\s*ИНН|\s*;\s*ОГРН|$)',
                        c,
                    )
                    if m:
                        name = m.group(1).strip().rstrip(";").strip()
                        # Пропускаем подсказки/шаблоны
                        if name and "полное наименование" not in name.lower() and len(name) > 3:
                            result["org_name"] = name
                            break

            # ИНН
            if "инн" in row_text:
                for c in cells:
                    m = re.search(r"\d{10,12}", c)
                    if m and "inn" not in result:
                        result["inn"] = m.group(0)

            # ОГРН
            if "огрн" in row_text:
                for c in cells:
                    m = re.search(r"\d{13,15}", c)
                    if m and "ogrn" not in result:
                        result["ogrn"] = m.group(0)

            # Адрес организации — «Почтовый адрес: …» в той же ячейке
            if "org_address" not in result and "почтовый адрес" in row_text:
                for c in cells:
                    m = re.search(r'[Пп]очтовый\s+адрес\s*:\s*([^;]+)', c)
                    if m:
                        addr = m.group(1).strip().rstrip(';').strip()
                        if addr and len(addr) > 5:
                            result["org_address"] = addr
                            break

            # Паспорт — объединяем: Наименование документа + Серия, номер + Дата выдачи + Кем выдан
            if "person_passport" not in result and (
                "наименование документа" in row_text or "серия" in row_text
                or "паспорт" in row_text or "документ, удостоверяющий" in row_text
            ):
                for c in cells:
                    m_doc  = re.search(r'[Нн]аименование\s+документа\s*:\s*([^;]+)', c)
                    m_ser  = re.search(r'[Сс]ерия[,\s]+номер\s*:\s*([^;]+)', c)
                    m_date = re.search(r'[Дд]ата\s+выдачи\s*:\s*([^;]+)', c)
                    m_by   = re.search(r'[Кк]ем\s+выдан\s*:\s*([^;]+)', c)
                    parts = [
                        g.group(1).strip().rstrip(';').strip()
                        for g in (m_doc, m_ser, m_date, m_by) if g
                    ]
                    if parts:
                        result["person_passport"] = ', '.join(p for p in parts if p)
                        break

            # Адрес физлица — «Адрес регистрации: …»
            if "person_address" not in result and "адрес регистрации" in row_text:
                for c in cells:
                    m = re.search(r'[Аа]дрес\s+регистрации\s*:\s*([^;]+)', c)
                    if m:
                        addr = m.group(1).strip().rstrip(';').strip()
                        if addr and len(addr) > 5:
                            result["person_address"] = addr
                            break

            # Срок размещения
            if "срок" in row_text and ("размещ" in row_text or "месяц" in row_text):
                for c in cells:
                    m = re.search(r"(\d+)\s*(месяц|мес)", c.lower())
                    if m and "term_months" not in result:
                        result["term_months"] = int(m.group(1))

    return result


def _extract_from_footers(doc: Document, result: dict) -> None:
    """
    Извлекает идентификатор ЕПГУ и дату заявления из колонтитулов.
    """
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer):
            if footer is None:
                continue
            for para in footer.paragraphs:
                text = (para.text or "").strip()
                if not text:
                    continue

                # Идентификатор на ЕПГУ → app_number
                if "app_number" not in result:
                    m = re.search(r'[Ии]дентификатор\s+на\s+ЕПГУ\s*:\s*(\d+)', text)
                    if m:
                        result["app_number"] = m.group(1)

                # Дата заявления → app_date
                if "app_date" not in result:
                    m = re.search(r'[Дд]ата\s+заявления\s*:\s*(\d{2})\.(\d{2})\.(\d{4})', text)
                    if m:
                        day, month, year = int(m.group(1)), int(m.group(2)), int(m.group(3))
                        result["app_date"] = date(year, month, day)
                        result["app_date_text"] = f"{day:02d}.{month:02d}.{year}"


def _extract_from_paragraphs(doc: Document, result: dict) -> None:
    """
    Извлекает данные из параграфов документа (срок размещения и т.д.).
    """
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        text_lower = text.lower()

        # Срок размещения из параграфа
        if "term_months" not in result and "срок" in text_lower:
            m = re.search(r'(?:на\s+срок|срок\s+размещения)[^:]*:\s*(\d+)\s*\(?месяц', text_lower)
            if m:
                result["term_months"] = int(m.group(1))


def parse_rrr_application_docx(doc_bytes: bytes) -> RRRApplicationData:
    """
    Парсер заявления РРР из DOCX файла.
    Извлекает данные заявителя, реквизиты заявления и срок действия.
    """
    doc = _load_doc(doc_bytes)
    data = _extract_from_tables(doc)
    _extract_from_footers(doc, data)
    _extract_from_paragraphs(doc, data)

    return RRRApplicationData(
        org_name=data.get("org_name"),
        inn=data.get("inn"),
        ogrn=data.get("ogrn"),
        org_address=data.get("org_address"),
        person_name=data.get("person_name"),
        person_passport=data.get("person_passport"),
        person_address=data.get("person_address"),
        app_number=data.get("app_number"),
        app_date=data.get("app_date"),
        app_date_text=data.get("app_date_text"),
        term_months=data.get("term_months"),
    )
