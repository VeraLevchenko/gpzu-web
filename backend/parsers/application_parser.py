# parsers/application_parser.py
"""
ОБНОВЛЁННАЯ ВЕРСИЯ с поддержкой телефона и email

ИЗМЕНЕНИЯ:
- Добавлено поле phone: Optional[str] = None
- Добавлено поле email: Optional[str] = None
- Добавлена функция _extract_phone_and_email_from_paragraphs()
- Обратная совместимость: старые модули не сломаются (поля опциональные)
"""
from __future__ import annotations

from dataclasses import dataclass
from datetime import date, timedelta
from io import BytesIO
from typing import Optional, Tuple
import re
from docx import Document


@dataclass
class ApplicationData:
    """
    Результат разбора заявления.
    
    ОБНОВЛЕНО: Добавлены поля phone и email
    """
    number: Optional[str] = None          # номер заявления
    date: Optional[date] = None           # дата заявления (объект date)
    date_text: Optional[str] = None       # строка даты, как в документе
    applicant: Optional[str] = None       # заявитель (ФИО или наименование ЮЛ)
    cadnum: Optional[str] = None          # кадастровый номер ЗУ
    purpose: Optional[str] = None         # цель использования ЗУ
    service_date: Optional[date] = None   # дата оказания услуги (14 раб. дней)
    
    # === НОВЫЕ ПОЛЯ === #
    phone: Optional[str] = None           # номер телефона заявителя
    email: Optional[str] = None           # адрес электронной почты заявителя


# ------------------------- ВСПОМОГАТЕЛЬНЫЕ ------------------------- #

def add_working_days(start_date: date, days: int = 14) -> date:
    """
    Добавляет N рабочих дней к дате, считая с самой даты (если она рабочая).
    Выходные: суббота/воскресенье. Праздники не учитываем.
    """
    d = start_date
    count = 0
    while True:
        if d.weekday() < 5:  # 0–4 = пн–пт
            count += 1
            if count >= days:
                return d
        d = d + timedelta(days=1)


def _load_doc(doc_bytes: bytes) -> Document:
    return Document(BytesIO(doc_bytes))


# ---------------------- ИЗВЛЕЧЕНИЕ ИЗ ТАБЛИЦ ---------------------- #

def _extract_number_and_date_from_tables(doc: Document) -> Tuple[Optional[str], Optional[date], Optional[str]]:
    """
    Номер и дата заявления берутся из первой таблицы, как в твоём шаблоне:
      таблица 0, строка 0:
        ячейка 0 -> '№: 6422028095'
        ячейка 1 -> '«15» ноября 2025 г.'
    """
    months = {
        "января": 1, "февраля": 2, "марта": 3, "апреля": 4, "мая": 5,
        "июня": 6, "июля": 7, "августа": 8, "сентября": 9, "октября": 10,
        "ноября": 11, "декабря": 12,
    }

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if not cells:
                continue

            # Ищем номер по шаблону "№: 6422028095"
            number = None
            for c in cells:
                m = re.search(r"[№N]\s*:?\s*([0-9]{5,})", c)
                if m:
                    number = m.group(1)
                    break

            # И ищем дату в виде «15» ноября 2025 г.
            d_val = None
            d_text = None
            for c in cells:
                if "«" in c and "»" in c and "г" in c:
                    d_text = c.strip()
                    try:
                        day_part = d_text.split("«", 1)[1].split("»", 1)[0]
                        day = int(day_part.strip().strip(". "))
                    except Exception:
                        day = None

                    rest = d_text.split("»", 1)[1]
                    rest = rest.replace("г.", "").replace("г", "").strip()
                    parts = rest.split()
                    month = None
                    year = None
                    if len(parts) >= 2:
                        month = months.get(parts[0].lower())
                        try:
                            year = int(parts[1])
                        except Exception:
                            year = None

                    if day and month and year:
                        try:
                            d_val = date(year, month, day)
                        except Exception:
                            d_val = None
                    break

            if number or d_val:
                return number, d_val, d_text

    return None, None, None


def _extract_applicant_from_tables(doc: Document) -> Optional[str]:
    """
    Заявитель:
      - если заполнено 1.2.1 (юр. лицо) — берём его,
      - иначе берём 1.1.1 (ФИО физлица).
    Оба блока в таблице 1 (по твоим примерам).
    """
    fio = None
    org = None

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if not cells:
                continue

            # 1.2.1 Полное наименование (юрлицо)
            if cells[0].startswith("1.2.1"):
                # обычно: [ '1.2.1', 'Полное наименование', '<Название ЮЛ>' ]
                if len(cells) >= 3 and cells[2]:
                    org = cells[2].strip()

            # 1.1.1 Фамилия, имя, отчество (ФЛ)
            if cells[0].startswith("1.1.1"):
                # обычно: [ '1.1.1', 'Фамилия, имя, отчество (при наличии)', 'ФИО' ]
                if len(cells) >= 3 and cells[2]:
                    fio = cells[2].strip()

    return org or fio


def _extract_cadnum_and_purpose_from_tables(doc: Document) -> Tuple[Optional[str], Optional[str]]:
    """
    Таблица 2 по шаблону:
      2.1 | 'Кадастровый номер земельного участка' | '42:...'
      2.3 | 'Цель использования земельного участка' | '...'
    """
    cad = None
    purpose = None

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if not cells:
                continue

            row_text = " | ".join(cells)

            # Кадастровый номер
            if "Кадастровый номер земельного участка" in row_text and cad is None:
                # обычно в третьей ячейке
                if len(cells) >= 3 and cells[2]:
                    # попытаемся вытащить реальный КН по маске
                    m = re.search(r"\d{2}:\d{2}:\d{6,7}:\d+", cells[2])
                    cad = m.group(0) if m else cells[2].strip()
                else:
                    # запасной вариант — берем последнюю непустую
                    for c in reversed(cells):
                        if c:
                            cad = c.strip()
                            break

            # Цель использования ЗУ
            if "Цель использования земельного участка" in row_text and purpose is None:
                if len(cells) >= 3 and cells[2]:
                    purpose = cells[2].strip()
                else:
                    for c in reversed(cells):
                        if c:
                            purpose = c.strip()
                            break

    return cad, purpose


# ---------------------- НОВАЯ ФУНКЦИЯ: ТЕЛЕФОН И EMAIL ---------------------- #

def _extract_phone_and_email_from_paragraphs(doc: Document) -> Tuple[Optional[str], Optional[str]]:
    """
    НОВОЕ: Извлекает телефон и email из параграфов документа.
    
    Ищет строку вида:
    "Номер телефона и адрес электронной почты для связи: +7(999)6485654 nazarasurov596@gmail.com"
    
    Returns:
        Tuple[phone, email]
    """
    phone = None
    email = None
    
    # Собираем весь текст из параграфов
    all_text = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text += text + "\n"
    
    # Ищем email (более специфичный паттерн, чтобы не захватить лишнее)
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    email_matches = re.findall(email_pattern, all_text)
    if email_matches:
        email = email_matches[0]  # Берём первый найденный email
    
    # Ищем телефон (10-11 цифр с возможными разделителями)
    # Паттерн: +7(...) или 8... или просто цифры
    phone_pattern = r'[\+\d][\d\s\-\(\)]{9,}'
    phone_matches = re.findall(phone_pattern, all_text)
    
    # Фильтруем: оставляем только те, где есть минимум 10 цифр
    valid_phones = []
    for p in phone_matches:
        digits = re.findall(r'\d', p)
        if len(digits) >= 10:
            valid_phones.append(p.strip())
    
    # Берём самый длинный телефон (обычно это контактный, а не из паспорта)
    if valid_phones:
        # Сортируем по длине и берём первый
        valid_phones.sort(key=lambda x: len(x), reverse=True)
        
        # Ищем телефон который начинается с +7 или 8 (российский формат)
        for p in valid_phones:
            if p.startswith('+7') or p.startswith('8'):
                phone = p
                break
        
        # Если не нашли российский - берём первый
        if not phone and valid_phones:
            phone = valid_phones[0]
    
    return phone, email


# ------------------------ ГЛАВНАЯ ФУНКЦИЯ ------------------------ #

def parse_application_docx(doc_bytes: bytes) -> ApplicationData:
    """
    Парсер заявления, заточенный под типовую форму "Заявление о выдаче ГПЗУ"
    (как в присланных примерах: всё основное в таблицах).
    
    ОБНОВЛЕНО: Теперь извлекает телефон и email из параграфов
    """
    doc = _load_doc(doc_bytes)

    # Извлекаем данные из таблиц (КАК РАНЬШЕ)
    number, app_date, date_text = _extract_number_and_date_from_tables(doc)
    applicant = _extract_applicant_from_tables(doc)
    cadnum, purpose = _extract_cadnum_and_purpose_from_tables(doc)

    # === НОВОЕ: Извлекаем телефон и email === #
    phone, email = _extract_phone_and_email_from_paragraphs(doc)

    # Рассчитываем дату оказания услуги
    service_date: Optional[date] = None
    if app_date:
        service_date = add_working_days(app_date, days=14)

    return ApplicationData(
        number=number,
        date=app_date,
        date_text=date_text,
        applicant=applicant,
        cadnum=cadnum,
        purpose=purpose,
        service_date=service_date,
        # === НОВЫЕ ПОЛЯ === #
        phone=phone,
        email=email,
    )