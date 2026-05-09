from io import BytesIO
from pathlib import Path

from docx import Document

TEMPLATE_PATH = (
    Path(__file__).parent.parent.parent / "templates" / "land_passport_template.docx"
)

# Порядок полей в шаблоне: строка i → название поля паспорта (i=1..28)
_FIELD_NAMES = [
    None,  # индекс 0 — заголовок таблицы
    "Кадастровый номер",
    "Кадастровый квартал",
    "Субъект РФ",
    "Муниципальное образование",
    "Адрес",
    "Площадь",
    "Категория земель",
    "Территориальная зона",
    "Вид разрешённого использования",
    "Форма собственности",
    "Кадастровая стоимость",
    "Потенциал использования",
    "Наличие объектов кап. строительства",
    "Наличие инженерных сетей",
    "Комплексное развитие территории",
    "Для льготных категорий граждан",
    "Льготная категория",
    "Агент АО «ДОМ.РФ»",
    "Расстояние до федеральной трассы",
    "Расстояние до дороги с твёрдым покрытием",
    "Расстояние до центра МО или ГО",
    "Расстояние до ближайшего населённого пункта",
    "Инвестиционный портал региона",
    "Наименование уполномоченного органа и его контакты",
    "Вовлечён под жилищное строительство",
    "Выдан ГПЗУ",
    "Выдано разрешение на строительство",
    "Отсутствует разрешение на ввод в эксплуатацию",
]


def _set_cell_text(table, row_idx: int, value: str) -> None:
    cell = table.rows[row_idx].cells[2]
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = value
        for extra in para.runs[1:]:
            extra.text = ""
    else:
        para.add_run(value)
    for extra_para in cell.paragraphs[1:]:
        for r in extra_para.runs:
            r.text = ""


def fill_passport(row_data: dict) -> bytes:
    """
    row_data — dict с ключами = названия полей из _FIELD_NAMES.
    Возвращает bytes заполненного docx.
    """
    doc = Document(TEMPLATE_PATH)
    table = doc.tables[0]

    for row_idx, field_name in enumerate(_FIELD_NAMES):
        if field_name is None:
            continue
        value = str(row_data.get(field_name, "") or "")
        _set_cell_text(table, row_idx, value)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
